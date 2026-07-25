"""
SchoolSheet Studio — Professional Worksheet Generator
Auto-assembled single-file app for deployment (e.g. Hugging Face Spaces).

This file was generated from the SchoolSheet_Studio notebook. All logic is
identical to the notebook version — only the packaging changed:
  - No Colab #@param widgets — AI provider + keys are read from environment
    variables / Space "Secrets" instead (set these in your hosting platform).
  - No !pip install lines — see requirements.txt instead.
  - App launches directly when this file runs (`python app.py`).
"""
import os
import requests

# ------------------------------------------------------------------
# AI Provider & Keys — configured via environment variables / Secrets
# (in Hugging Face Spaces: Settings -> Variables and secrets).
#   AI_PROVIDER          = "anthropic" | "gemini" | "none"   (optional; auto-detected if unset)
#   ANTHROPIC_API_KEY     = your Anthropic key   (if using Claude)
#   GEMINI_API_KEY         = your Gemini key      (if using Gemini)
#   PIXABAY_API_KEY        = optional, for real photo illustrations
# ------------------------------------------------------------------
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "").strip()
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
PIXABAY_API_KEY = os.environ.get("PIXABAY_API_KEY", "").strip()
_provider_pref = os.environ.get("AI_PROVIDER", "").strip().lower()

if _provider_pref == "anthropic" and ANTHROPIC_API_KEY:
    AI_ENABLED = True
    ACTIVE_PROVIDER = "anthropic"
elif _provider_pref == "gemini" and GEMINI_API_KEY:
    AI_ENABLED = True
    ACTIVE_PROVIDER = "gemini"
elif _provider_pref in ("none", ""):
    # Auto-detect when AI_PROVIDER isn't explicitly set: prefer Anthropic,
    # then Gemini, else disabled.
    if ANTHROPIC_API_KEY:
        AI_ENABLED = True
        ACTIVE_PROVIDER = "anthropic"
    elif GEMINI_API_KEY:
        AI_ENABLED = True
        ACTIVE_PROVIDER = "gemini"
    else:
        AI_ENABLED = False
        ACTIVE_PROVIDER = None
else:
    AI_ENABLED = False
    ACTIVE_PROVIDER = None

if AI_ENABLED:
    _provider_label = "Anthropic (Claude)" if ACTIVE_PROVIDER == "anthropic" else "Google (Gemini)"
    print(f"AI provider active: {_provider_label}")
else:
    print("No AI provider configured — running in Question Bank-only mode.")

# ------------------------------------------------------------------
# Urdu Unicode font (Noto Nastaliq Urdu) so Urdu text renders correctly
# in PDF and Word output.
# ------------------------------------------------------------------
FONT_DIR = "./fonts"
os.makedirs(FONT_DIR, exist_ok=True)
URDU_FONT_PATH = os.path.join(FONT_DIR, "NotoNastaliqUrdu-Regular.ttf")

if not os.path.exists(URDU_FONT_PATH):
    font_url = "https://raw.githubusercontent.com/google/fonts/main/ofl/notonastaliqurdu/NotoNastaliqUrdu%5Bwght%5D.ttf"
    try:
        r = requests.get(font_url, timeout=30)
        r.raise_for_status()
        with open(URDU_FONT_PATH, "wb") as f:
            f.write(r.content)
        print("Urdu font downloaded successfully.")
    except Exception as e:
        print("Could not auto-download Urdu font, will fall back to a bundled system font.", e)
        URDU_FONT_PATH = None

LATIN_FONT_PATH = None
try:
    import matplotlib
    LATIN_FONT_PATH = os.path.join(os.path.dirname(matplotlib.__file__), "mpl-data", "fonts", "ttf", "DejaVuSans.ttf")
    if not os.path.exists(LATIN_FONT_PATH):
        LATIN_FONT_PATH = None
except Exception:
    LATIN_FONT_PATH = None

print("Setup complete. Urdu font path:", URDU_FONT_PATH)

import pdfplumber
import docx as docx_lib

# ------------------------------------------------------------------
# Unified AI call layer — one function, works with EITHER provider.
# Every other step in this notebook calls _call_ai() and never talks
# to the Anthropic or Gemini SDKs directly, so switching providers in
# Step 2 doesn't require touching any code below.
# ------------------------------------------------------------------
_anthropic_client = None
_gemini_model = None

CLAUDE_MODEL = "claude-sonnet-4-6"
GEMINI_MODEL = "gemini-2.5-flash"

if AI_ENABLED and ACTIVE_PROVIDER == "anthropic":
    from anthropic import Anthropic
    _anthropic_client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
elif AI_ENABLED and ACTIVE_PROVIDER == "gemini":
    import google.generativeai as genai
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    _gemini_model = genai.GenerativeModel(GEMINI_MODEL)


def _call_ai(system_prompt: str, user_content: str, max_tokens: int = 4000) -> str:
    """
    Single entry point for every AI call in this notebook.
    Sends `user_content` with `system_prompt` as instructions and returns
    the raw text reply — regardless of which provider is active.
    Raises RuntimeError if AI is not enabled (no valid key/provider).
    """
    if not AI_ENABLED:
        raise RuntimeError(
            "AI generation is disabled because no valid API key was entered in Step 2. "
            "Use the 'Question Bank' method in the dashboard instead, or add a key in Step 2 and re-run it."
        )
    if ACTIVE_PROVIDER == "anthropic":
        resp = _anthropic_client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=max_tokens,
            system=system_prompt,
            messages=[{"role": "user", "content": user_content}],
        )
        return "".join(b.text for b in resp.content if b.type == "text").strip()
    elif ACTIVE_PROVIDER == "gemini":
        full_prompt = f"{system_prompt}\n\n{user_content}"
        resp = _gemini_model.generate_content(
            full_prompt,
            generation_config={"max_output_tokens": max_tokens},
        )
        return (getattr(resp, "text", None) or "").strip()
    else:
        raise RuntimeError(f"Unknown AI provider: {ACTIVE_PROVIDER}")


def extract_text_from_file(file_path: str) -> str:
    """Reads .pdf, .docx or .txt and returns raw text."""
    if file_path is None:
        return ""
    lower = file_path.lower()
    try:
        if lower.endswith(".pdf"):
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    text_parts.append(t)
            return "\n".join(text_parts)
        elif lower.endswith(".docx"):
            d = docx_lib.Document(file_path)
            return "\n".join(p.text for p in d.paragraphs)
        else:  # .txt or anything else — read as plain text
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        return f"[ERROR READING FILE: {e}]"


def clean_pasted_content(raw_text: str) -> str:
    """
    Removes AI chatter / preambles / postambles / instructions that are
    NOT actual question content (e.g. "Here are your questions:",
    "I hope this helps!", numbering artifacts from chat exports, etc.)
    Keeps only the real educational content, unaltered.
    """
    if not raw_text or not raw_text.strip():
        return ""

    if not AI_ENABLED:
        # Simple rule-based fallback cleaning when no API key is configured:
        # drops common AI-chatter lines and blank/short filler lines.
        junk_markers = [
            "here are your questions", "here's your worksheet", "i hope this helps",
            "hope this helps", "let me know if", "feel free to", "sure, here",
            "certainly!", "as requested", "best regards", "regards,",
        ]
        lines = raw_text.splitlines()
        kept = []
        for line in lines:
            l = line.strip()
            if not l:
                continue
            if any(marker in l.lower() for marker in junk_markers):
                continue
            kept.append(line)
        return "\n".join(kept).strip()

    system_prompt = (
        "You are a strict text-cleaning engine for a school worksheet system. "
        "You will be given raw text that may contain educational questions or "
        "content mixed with irrelevant chatter (greetings, meta-comments like "
        "'Here are your questions', 'I hope this helps', signatures, extra "
        "headers, disclaimers, or duplicate/broken lines). "
        "Return ONLY the genuine educational content — questions, passages, "
        "topics, instructions meant for students — exactly as originally "
        "written, with no rephrasing. Remove every non-content line. "
        "Do not add any commentary, notes, or explanation of your own. "
        "Output nothing except the cleaned content."
    )

    return _call_ai(system_prompt, raw_text, max_tokens=4000)


print("Text extraction & cleaning utilities ready. Active AI provider:", ACTIVE_PROVIDER or "None (Question Bank mode)")

from dataclasses import dataclass, field
from typing import List, Optional
import json, random

QUESTION_TYPES = [
    "Choose the Best Option (MCQ)",
    "Fill in the Blanks",
    "Match the Column",
    "Short Question",
    "Detailed Question",
]

GRADES = [str(g) for g in range(1, 11)]  # Grade 1 to Grade 10

THEMES = ["Auto (based on grade)", "Colorful / Playful", "Professional / Formal"]


@dataclass
class QuestionItem:
    q_type: str
    text: str
    options: Optional[List[str]] = None       # for MCQ
    match_left: Optional[List[str]] = None     # for Match the Column
    match_right: Optional[List[str]] = None
    answer: Optional[str] = None
    image_keyword: Optional[str] = None        # e.g. "apple", "cat" for young grades

    def to_dict(self):
        return self.__dict__


@dataclass
class WorksheetConfig:
    school_name: str = "Your School Name"
    logo_path: Optional[str] = None
    logo_size_pt: int = 60          # logo width in points
    cover_color_hex: str = "#2563EB"
    font_theme: str = "Poppins/Default"

    grade: str = "3"
    subject: str = "General"
    topic: str = "General Topic"
    language: str = "English"       # "English" or "Urdu"

    question_types: List[str] = field(default_factory=lambda: ["Choose the Best Option (MCQ)"])
    num_questions: int = 10

    questions_per_page: Optional[int] = None   # None = unlimited/auto-flow
    spacing_pt: int = 18                        # vertical space between questions
    binding_gap_mm: int = 0                      # extra margin for binding
    num_copies: int = 1
    copy_mode: str = "Identical"                 # "Identical" or "Randomized"

    visual_theme: str = "Auto (based on grade)"

    show_student_name: bool = True
    show_class: bool = True
    show_roll_no: bool = True
    show_date: bool = True

    def effective_theme(self) -> str:
        if self.visual_theme != "Auto (based on grade)":
            return self.visual_theme
        return "Colorful / Playful" if int(self.grade) <= 4 else "Professional / Formal"


print("Data models ready. Question types:", QUESTION_TYPES)

import pandas as pd

def _build_generation_prompt(cfg: "WorksheetConfig", source_text: Optional[str]) -> str:
    types_str = ", ".join(cfg.question_types)
    lang_line = (
        "Write the ENTIRE worksheet content in proper, correct, natural Urdu "
        "(Urdu script, not Roman Urdu)."
        if cfg.language == "Urdu"
        else "Write everything in clear, simple, grade-appropriate English."
    )
    source_block = (
        f"Base every question strictly on the following source content "
        f"(do not invent facts outside it):\n\n{source_text}\n"
        if source_text
        else f"Topic: {cfg.topic}\nSubject: {cfg.subject}\n"
    )
    image_line = (
        "Since this is for a young grade (1-4), for at least half the "
        "questions add a simple, common, real-world 'image_keyword' "
        "(e.g. apple, cat, monkey, ball, sun, book) that a teacher could "
        "illustrate the question with. Use only simple everyday objects/animals."
        if int(cfg.grade) <= 4
        else "This is for an older grade — do not add image_keyword, keep it null."
    )

    return f"""
You are an expert school curriculum question-writer for Grade {cfg.grade}, subject {cfg.subject}.
{source_block}
Generate exactly {cfg.num_questions} TOTAL questions, using ONLY these question types: {types_str}.
Distribute the questions across the selected types as evenly as possible.
{lang_line}
{image_line}

Rules:
- Every question must be completely different in content from every other question (no repeats, no rephrasing of the same fact twice).
- For "Choose the Best Option (MCQ)": provide exactly 4 options and mark the correct answer.
- For "Fill in the Blanks": use a blank shown as "_____" inside the sentence, and give the answer.
- For "Match the Column": provide a left list and right list (same length, 4-6 pairs), shuffle the right list order, and give the answer as pairs.
- For "Short Question": a question needing a 1-3 sentence answer.
- For "Detailed Question": a question needing a paragraph-length answer.
- Keep difficulty appropriate for Grade {cfg.grade}.

Return ONLY valid JSON (no markdown fences, no commentary), as a list of objects with this exact schema:
[
  {{
    "q_type": "one of: Choose the Best Option (MCQ) | Fill in the Blanks | Match the Column | Short Question | Detailed Question",
    "text": "the question text (or sentence with blank)",
    "options": ["...", "...", "...", "..."] or null,
    "match_left": ["..."] or null,
    "match_right": ["..."] or null,
    "answer": "the correct answer / answer key",
    "image_keyword": "simple keyword" or null
  }}
]
""".strip()


def generate_questions_ai(cfg: "WorksheetConfig", source_text: Optional[str] = None) -> List[QuestionItem]:
    """Calls the active AI provider (Anthropic or Gemini — see Step 2) to generate a fresh, non-repeating question set."""
    if not AI_ENABLED:
        raise RuntimeError(
            "AI generation is disabled because no valid API key was entered in Step 2. "
            "Use the 'Question Bank' method in the dashboard instead, or add a key in Step 2 and re-run it."
        )
    prompt = _build_generation_prompt(cfg, source_text)
    raw = _call_ai(
        system_prompt=(
            "You are an expert school curriculum question-writer. "
            "Return ONLY valid JSON (no markdown fences, no commentary)."
        ),
        user_content=prompt,
        max_tokens=4000,
    )
    raw = raw.strip("`")
    if raw.lower().startswith("json"):
        raw = raw[4:].strip()
    try:
        data = json.loads(raw)
    except Exception:
        # last-resort: try to locate the JSON array inside the text
        start = raw.find("[")
        end = raw.rfind("]")
        data = json.loads(raw[start:end + 1])

    return [QuestionItem(
        q_type=item.get("q_type", "Short Question"),
        text=item.get("text", ""),
        options=item.get("options"),
        match_left=item.get("match_left"),
        match_right=item.get("match_right"),
        answer=item.get("answer"),
        image_keyword=item.get("image_keyword"),
    ) for item in data]


def load_question_bank(file_path: str) -> pd.DataFrame:
    """
    Optional plug-in question bank. Expected columns (csv or xlsx):
    grade, subject, topic, q_type, text, options, match_left, match_right, answer, image_keyword
    'options' / 'match_left' / 'match_right' cells should be pipe-separated, e.g.  "A|B|C|D"
    """
    if file_path.lower().endswith(".csv"):
        return pd.read_csv(file_path)
    return pd.read_excel(file_path)


def questions_from_bank(df: "pd.DataFrame", cfg: "WorksheetConfig") -> List[QuestionItem]:
    filtered = df[
        (df["grade"].astype(str) == str(cfg.grade))
        & (df["subject"].str.lower() == cfg.subject.lower())
        & (df["topic"].str.lower() == cfg.topic.lower())
        & (df["q_type"].isin(cfg.question_types))
    ]
    filtered = filtered.sample(frac=1).head(cfg.num_questions)  # shuffle + limit
    items = []
    for _, row in filtered.iterrows():
        split = lambda v: str(v).split("|") if pd.notna(v) else None
        items.append(QuestionItem(
            q_type=row["q_type"],
            text=row["text"],
            options=split(row.get("options")),
            match_left=split(row.get("match_left")),
            match_right=split(row.get("match_right")),
            answer=row.get("answer"),
            image_keyword=row.get("image_keyword") if pd.notna(row.get("image_keyword")) else None,
        ))
    return items


def randomize_copy(questions: List[QuestionItem]) -> List[QuestionItem]:
    """Produces a shuffled variant (question order + MCQ option order) for randomized copies."""
    import copy as _copy
    shuffled = _copy.deepcopy(questions)
    random.shuffle(shuffled)
    for q in shuffled:
        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            correct = q.options[0] if q.answer == q.options[0] else q.answer
            opts = q.options[:]
            random.shuffle(opts)
            q.options = opts
    return shuffled


print("AI generation engine + question bank plug-in ready.")

import hashlib

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_OK = True
except Exception as e:
    print("⚠️ Pillow (PIL) failed to import:", e)
    print("   Images will be skipped, but the rest of the app (PDF/Word generation,")
    print("   AI question generation, all worksheet features) will still work fine.")
    print("   To fix images: Runtime -> Restart session, then run all cells again from Step 1.")
    PIL_OK = False

IMAGE_CACHE_DIR = "./wg_images"
os.makedirs(IMAGE_CACHE_DIR, exist_ok=True)

_FALLBACK_COLORS = [
    "#FF6B6B", "#4ECDC4", "#FFD93D", "#6BCB77", "#4D96FF", "#FF922B", "#C084FC"
]


def _cache_path(keyword: str) -> str:
    h = hashlib.md5(keyword.encode()).hexdigest()[:10]
    return os.path.join(IMAGE_CACHE_DIR, f"{h}.png")


def _make_fallback_icon(keyword: str, path: str, size: int = 300):
    """Generates a clean, simple flat-color icon with the keyword label,
    used only if no Pixabay key / no internet result is available."""
    if not PIL_OK:
        return None
    color = _FALLBACK_COLORS[hash(keyword) % len(_FALLBACK_COLORS)]
    img = Image.new("RGB", (size, size), "white")
    draw = ImageDraw.Draw(img)
    draw.ellipse([20, 20, size - 20, size - 20], fill=color)
    try:
        font = ImageFont.truetype(LATIN_FONT_PATH, 28) if LATIN_FONT_PATH else ImageFont.load_default()
    except Exception:
        font = ImageFont.load_default()
    text = keyword.upper()
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((size - tw) / 2, size - 55), text, fill="white", font=font)
    img.save(path)
    return path


def get_image_for_keyword(keyword: str) -> Optional[str]:
    """Returns a local file path to an image representing `keyword`.
    Tries Pixabay (real photo) first if a key is set, else falls back
    to a clean generated icon so the worksheet is never left blank.
    Returns None (image simply skipped) if Pillow is unavailable."""
    if not keyword:
        return None
    path = _cache_path(keyword)
    if os.path.exists(path):
        return path

    pixabay_key = os.environ.get("PIXABAY_API_KEY", "").strip()
    if pixabay_key:
        try:
            r = requests.get(
                "https://pixabay.com/api/",
                params={
                    "key": pixabay_key,
                    "q": keyword,
                    "image_type": "photo",
                    "category": "animals,food,education",
                    "safesearch": "true",
                    "per_page": 3,
                },
                timeout=15,
            )
            data = r.json()
            if data.get("hits"):
                img_url = data["hits"][0]["webformatURL"]
                img_data = requests.get(img_url, timeout=15).content
                with open(path, "wb") as f:
                    f.write(img_data)
                return path
        except Exception as e:
            print(f"Pixabay fetch failed for '{keyword}', using fallback icon. ({e})")

    return _make_fallback_icon(keyword, path)


print("Image engine ready (real photos via Pixabay, with automatic icon fallback).")

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display

# Register fonts once
_FONTS_REGISTERED = False

def _register_fonts():
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    if URDU_FONT_PATH and os.path.exists(URDU_FONT_PATH):
        pdfmetrics.registerFont(TTFont("Urdu", URDU_FONT_PATH))
    if LATIN_FONT_PATH and os.path.exists(LATIN_FONT_PATH):
        pdfmetrics.registerFont(TTFont("Latin", LATIN_FONT_PATH))
    _FONTS_REGISTERED = True

def _shape(text: str, is_urdu: bool) -> str:
    if not is_urdu:
        return text
    reshaped = arabic_reshaper.reshape(text)
    return get_display(reshaped)


THEME_COLORS = {
    "Colorful / Playful": {"primary": "#FF6B6B", "secondary": "#FFD93D", "accent": "#4ECDC4"},
    "Professional / Formal": {"primary": "#1E3A8A", "secondary": "#334155", "accent": "#2563EB"},
}


def render_pdf(cfg: "WorksheetConfig", questions: List[QuestionItem], output_path: str):
    _register_fonts()
    is_urdu = cfg.language == "Urdu"
    body_font = "Urdu" if (is_urdu and "Urdu" in pdfmetrics.getRegisteredFontNames()) else "Helvetica"
    head_font = "Latin" if "Latin" in pdfmetrics.getRegisteredFontNames() else "Helvetica-Bold"

    theme = THEME_COLORS[cfg.effective_theme()]
    page_w, page_h = A4
    left_margin = 18 * mm + cfg.binding_gap_mm * mm
    right_margin = 18 * mm
    top_margin = 15 * mm
    bottom_margin = 15 * mm

    c = canvas.Canvas(output_path, pagesize=A4)

    def draw_header(page_no_label=""):
        y = page_h - top_margin
        # Logo
        if cfg.logo_path and os.path.exists(cfg.logo_path):
            try:
                c.drawImage(cfg.logo_path, left_margin, y - cfg.logo_size_pt,
                            width=cfg.logo_size_pt, height=cfg.logo_size_pt,
                            preserveAspectRatio=True, mask="auto")
            except Exception:
                pass
        # School name
        c.setFont(head_font, 16)
        c.setFillColor(theme["primary"])
        c.drawCentredString(page_w / 2, y - 20, _shape(cfg.school_name, False))
        c.setFont(head_font, 11)
        c.setFillColor("#333333")
        c.drawCentredString(page_w / 2, y - 38, f"Grade {cfg.grade}  |  {cfg.subject}  |  {cfg.topic}")

        # divider line
        c.setStrokeColor(theme["accent"])
        c.setLineWidth(1.5)
        c.line(left_margin, y - 48, page_w - right_margin, y - 48)

        # student info fields
        fy = y - 65
        c.setFont("Helvetica", 10)
        c.setFillColor("#000000")
        fields = []
        if cfg.show_student_name: fields.append("Student Name: ______________________")
        if cfg.show_class: fields.append("Class: _______")
        if cfg.show_roll_no: fields.append("Roll No: _______")
        if cfg.show_date: fields.append("Date: _______")
        line1 = "     ".join(fields[:2])
        line2 = "     ".join(fields[2:])
        c.drawString(left_margin, fy, line1)
        if line2:
            c.drawString(left_margin, fy - 14, line2)

        c.setStrokeColor("#999999")
        c.setLineWidth(0.7)
        c.line(left_margin, fy - 24, page_w - right_margin, fy - 24)
        return fy - 40  # y position where question body should start

    def new_page():
        c.showPage()
        return draw_header()

    y_cursor = draw_header()
    q_count_on_page = 0
    max_width = page_w - left_margin - right_margin

    def wrap_text(text, font, size, max_w):
        c.setFont(font, size)
        words = text.split(" ")
        lines, cur = [], ""
        for w in words:
            trial = (cur + " " + w).strip()
            if c.stringWidth(trial, font, size) <= max_w:
                cur = trial
            else:
                lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
        return lines

    for idx, q in enumerate(questions, start=1):
        needed_space = 70  # rough estimate; grows below per type
        if y_cursor < bottom_margin + needed_space or (
            cfg.questions_per_page and q_count_on_page >= cfg.questions_per_page
        ):
            y_cursor = new_page()
            q_count_on_page = 0

        c.setFont(body_font, 11)
        c.setFillColor("#000000")
        prefix = f"{idx}. "
        q_text = _shape(prefix + q.text, is_urdu)
        lines = wrap_text(q_text, body_font, 11, max_width)
        for line in lines:
            c.drawString(left_margin, y_cursor, line)
            y_cursor -= 15

        # image for young grades
        if q.image_keyword:
            img_path = get_image_for_keyword(q.image_keyword)
            if img_path:
                try:
                    c.drawImage(img_path, left_margin, y_cursor - 60, width=55, height=55,
                                preserveAspectRatio=True, mask="auto")
                    y_cursor -= 65
                except Exception:
                    pass

        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            for i, opt in enumerate(q.options):
                letter = chr(97 + i)
                c.drawString(left_margin + 15, y_cursor, _shape(f"({letter}) {opt}", is_urdu))
                y_cursor -= 14

        elif q.q_type == "Match the Column" and q.match_left and q.match_right:
            col2_x = left_margin + max_width / 2
            rows = max(len(q.match_left), len(q.match_right))
            for i in range(rows):
                left_txt = q.match_left[i] if i < len(q.match_left) else ""
                right_txt = q.match_right[i] if i < len(q.match_right) else ""
                c.drawString(left_margin + 15, y_cursor, _shape(f"{i+1}. {left_txt}", is_urdu))
                c.drawString(col2_x, y_cursor, _shape(f"{chr(97+i)}. {right_txt}", is_urdu))
                y_cursor -= 14

        elif q.q_type == "Short Question":
            for _ in range(2):
                y_cursor -= 14
                c.setStrokeColor("#cccccc")
                c.line(left_margin, y_cursor, page_w - right_margin, y_cursor)

        elif q.q_type == "Detailed Question":
            for _ in range(5):
                y_cursor -= 14
                c.setStrokeColor("#cccccc")
                c.line(left_margin, y_cursor, page_w - right_margin, y_cursor)

        elif q.q_type == "Fill in the Blanks":
            pass  # blank already embedded in text

        y_cursor -= cfg.spacing_pt
        q_count_on_page += 1

    c.save()
    return output_path


def build_pdf_copies(cfg: "WorksheetConfig", base_questions: List[QuestionItem], output_dir: str) -> List[str]:
    os.makedirs(output_dir, exist_ok=True)
    paths = []
    for i in range(1, cfg.num_copies + 1):
        qs = randomize_copy(base_questions) if cfg.copy_mode == "Randomized" else base_questions
        out = os.path.join(output_dir, f"worksheet_copy_{i}.pdf")
        render_pdf(cfg, qs, out)
        paths.append(out)
    return paths


print("PDF renderer ready.")

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _apply_urdu_font(run, size=12):
    run.font.name = "Noto Nastaliq Urdu"
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), "Noto Nastaliq Urdu")


def render_docx(cfg: "WorksheetConfig", questions: List[QuestionItem], output_path: str):
    is_urdu = cfg.language == "Urdu"
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Mm(18 + cfg.binding_gap_mm)
    section.right_margin = Mm(18)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    theme_primary = RGBColor(0x1E, 0x3A, 0x8A) if cfg.effective_theme() == "Professional / Formal" else RGBColor(0xE0, 0x50, 0x50)

    # Logo + School name
    if cfg.logo_path and os.path.exists(cfg.logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(cfg.logo_path, width=Pt(cfg.logo_size_pt))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(cfg.school_name)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = theme_primary

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Grade {cfg.grade}  |  {cfg.subject}  |  {cfg.topic}").font.size = Pt(11)

    doc.add_paragraph("_" * 90)

    # Student fields
    fields = []
    if cfg.show_student_name: fields.append("Student Name: ______________________")
    if cfg.show_class: fields.append("Class: _______")
    if cfg.show_roll_no: fields.append("Roll No: _______")
    if cfg.show_date: fields.append("Date: _______")
    if fields:
        doc.add_paragraph("     ".join(fields[:2]))
        if len(fields) > 2:
            doc.add_paragraph("     ".join(fields[2:]))
    doc.add_paragraph("_" * 90)

    for idx, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        if is_urdu:
            _set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{idx}. {q.text}")
        run.bold = True
        if is_urdu:
            _apply_urdu_font(run, 13)
        else:
            run.font.size = Pt(12)

        if q.image_keyword:
            img_path = get_image_for_keyword(q.image_keyword)
            if img_path:
                try:
                    doc.add_picture(img_path, width=Inches(0.7))
                except Exception:
                    pass

        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            for i, opt in enumerate(q.options):
                op = doc.add_paragraph(f"    ({chr(97+i)}) {opt}")
                if is_urdu:
                    _set_rtl(op)
                    op.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _apply_urdu_font(op.runs[0], 12)

        elif q.q_type == "Match the Column" and q.match_left and q.match_right:
            table = doc.add_table(rows=max(len(q.match_left), len(q.match_right)), cols=2)
            table.style = "Table Grid"
            for i in range(len(table.rows)):
                l = q.match_left[i] if i < len(q.match_left) else ""
                r_ = q.match_right[i] if i < len(q.match_right) else ""
                table.cell(i, 0).text = f"{i+1}. {l}"
                table.cell(i, 1).text = f"{chr(97+i)}. {r_}"

        elif q.q_type == "Short Question":
            for _ in range(2):
                doc.add_paragraph("_" * 90)

        elif q.q_type == "Detailed Question":
            for _ in range(5):
                doc.add_paragraph("_" * 90)

        # spacing between questions
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(cfg.spacing_pt)

        if cfg.questions_per_page and idx % cfg.questions_per_page == 0 and idx != len(questions):
            doc.add_page_break()

    doc.save(output_path)
    return output_path


def build_docx_copies(cfg: "WorksheetConfig", base_questions: List[QuestionItem], output_dir: str) -> List[str]:
    os.makedirs(output_dir, exist_ok=True)
    paths = []
    for i in range(1, cfg.num_copies + 1):
        qs = randomize_copy(base_questions) if cfg.copy_mode == "Randomized" else base_questions
        out = os.path.join(output_dir, f"worksheet_copy_{i}.docx")
        render_docx(cfg, qs, out)
        paths.append(out)
    return paths


print("DOCX renderer ready.")

import gradio as gr
import tempfile, shutil

APP_NAME = "SchoolSheet Studio"  # Worksheet Generator name — change as you like
WORK_DIR = "./wg_output"
os.makedirs(WORK_DIR, exist_ok=True)

SUBJECTS = ["English", "Urdu", "Mathematics", "Science", "Social Studies", "Islamiyat", "Computer Science", "General Knowledge"]


def build_config(school_name, logo_file, logo_size, cover_color,
                  grade, subject, topic, language,
                  q_types, num_q,
                  q_per_page, spacing, binding_gap, num_copies, copy_mode,
                  visual_theme,
                  show_name, show_class, show_roll, show_date):
    return WorksheetConfig(
        school_name=school_name or "Your School Name",
        logo_path=logo_file.name if logo_file else None,
        logo_size_pt=int(logo_size),
        cover_color_hex=cover_color,
        grade=str(grade),
        subject=subject,
        topic=topic or "General Topic",
        language=language,
        question_types=q_types or ["Choose the Best Option (MCQ)"],
        num_questions=int(num_q),
        questions_per_page=(int(q_per_page) if q_per_page and int(q_per_page) > 0 else None),
        spacing_pt=int(spacing),
        binding_gap_mm=int(binding_gap),
        num_copies=int(num_copies),
        copy_mode=copy_mode,
        visual_theme=visual_theme,
        show_student_name=show_name, show_class=show_class,
        show_roll_no=show_roll, show_date=show_date,
    )


def preview_text(questions: List[QuestionItem]) -> str:
    if not questions:
        return "No questions generated yet."
    out = []
    for i, q in enumerate(questions, 1):
        out.append(f"**{i}. ({q.q_type})** {q.text}")
        if q.options:
            out.append("   " + "   ".join(f"({chr(97+j)}) {o}" for j, o in enumerate(q.options)))
        if q.match_left and q.match_right:
            out.append(f"   Match: {q.match_left}  ↔  {q.match_right}")
        if q.image_keyword:
            out.append(f"   🖼 illustration: {q.image_keyword}")
        out.append("")
    return "\n".join(out)


if AI_ENABLED:
    _provider_name = "Anthropic (Claude)" if ACTIVE_PROVIDER == "anthropic" else "Google (Gemini)"
    AI_STATUS_BANNER = f"✅ AI question generation is **enabled** using **{_provider_name}** (configured in Step 2)."
else:
    AI_STATUS_BANNER = (
        "ℹ️ AI question generation is **off** (no valid AI provider/key set in Step 2). "
        "Methods 1 & 2 below need it — use **Method 3: Question Bank** instead, which works with zero API key/cost."
    )

GETTING_STARTED_GUIDE = """
### 👋 Naya user? Yahan se shuru karein — New here? Start with this

**1. Worksheet Settings bharein:** Neeche "⚙️ Worksheet Settings" section mein Grade, Subject, Topic, language aur
question types chunein — ye saare 3 methods ke liye common hain.

**2. Questions banayein — 3 tareeqe (methods) mein se koi ek chunein:**
- 📂 **Method 1 — Curriculum-Based:** Apni file (PDF/DOCX/TXT) upload karein YA sirf Grade/Subject/Topic se generate karein. *(AI zaroori hai)*
- 📝 **Method 2 — Free Content:** Koi bhi text (jaise ChatGPT se copy kiya hua) paste ya upload karein — app khud saaf kar dega. *(AI zaroori hai)*
- 🗂️ **Method 3 — Apna Question Bank:** Apni CSV/Excel file upload karein — bilkul free, AI ki zaroorat nahi.

**3. Preview dekhein aur Export karein:** Neeche "👁️ Preview & Export" section mein PDF ya Word format chunein aur
**"📄 Create Worksheet"** button dabayein — download link turant mil jayega.

> 💡 **Tip:** Agar Methods 1 ya 2 kaam nahi kar rahe to iska matlab is app ke admin ne AI provider set nahi kiya —
> seedha Method 3 (Question Bank) use karein, ya admin se AI key add karwayein.
""".strip()


# ---------------- Method 1 : Path A (file + chapter/topic/page) ----------------
def generate_method1_pathA(file_obj, chapter, topic_a, page_no, *shared_args):
    if not file_obj:
        return "⚠️ Please upload a curriculum file first.", None
    try:
        raw = extract_text_from_file(file_obj.name)
        focus_note = f"\n\n[Focus strictly on: Chapter '{chapter}', Topic '{topic_a}', Page {page_no}]"
        cfg = build_config(*shared_args)
        cfg.topic = topic_a or cfg.topic
        questions = generate_questions_ai(cfg, source_text=raw + focus_note)
        return preview_text(questions), questions
    except RuntimeError as e:
        return f"⚠️ {e}", None


# ---------------- Method 1 : Path B (grade/subject/topic dropdown) ----------------
def generate_method1_pathB(*shared_args):
    try:
        cfg = build_config(*shared_args)
        questions = generate_questions_ai(cfg, source_text=None)
        return preview_text(questions), questions
    except RuntimeError as e:
        return f"⚠️ {e}", None


# ---------------- Method 2 : paste / upload any content ----------------
def generate_method2(pasted_text, file_obj, *shared_args):
    raw = pasted_text or ""
    if file_obj:
        raw += "\n" + extract_text_from_file(file_obj.name)
    if not raw.strip():
        return "⚠️ Please paste text or upload a file first.", None
    try:
        cleaned = clean_pasted_content(raw)
        cfg = build_config(*shared_args)
        questions = generate_questions_ai(cfg, source_text=cleaned)
        return preview_text(questions), questions
    except RuntimeError as e:
        return f"⚠️ {e}", None


# ---------------- Method 3 : your own question bank (no AI / no key needed) ----------------
def generate_method3(bank_file, *shared_args):
    if not bank_file:
        return "⚠️ Please upload a question bank CSV/Excel file first.", None
    try:
        cfg = build_config(*shared_args)
        df = load_question_bank(bank_file.name)
        questions = questions_from_bank(df, cfg)
        if not questions:
            return ("⚠️ No matching questions found in your bank for this Grade/Subject/Topic/Question-Type "
                     "combination. Check the values in Worksheet Settings above match your file's columns."), None
        return preview_text(questions), questions
    except Exception as e:
        return f"⚠️ Could not read question bank: {e}", None


# ---------------- Final export ----------------
def export_worksheet(questions_state, export_format, *shared_args):
    if not questions_state:
        return None, "⚠️ Please generate/preview a worksheet first."
    cfg = build_config(*shared_args)
    out_dir = tempfile.mkdtemp(dir=WORK_DIR)
    if export_format == "PDF":
        paths = build_pdf_copies(cfg, questions_state, out_dir)
    else:
        paths = build_docx_copies(cfg, questions_state, out_dir)

    if len(paths) == 1:
        return paths[0], f"✅ Worksheet generated ({export_format}) — {len(paths)} copy."
    else:
        zip_path = shutil.make_archive(os.path.join(out_dir, "worksheets"), "zip", out_dir)
        return zip_path, f"✅ {len(paths)} {export_format} copies generated and zipped."


# =====================================================================================
#  UI LAYOUT
# =====================================================================================
with gr.Blocks(title=APP_NAME, theme=gr.themes.Soft()) as app:
    gr.Markdown(f"# 📘 {APP_NAME}\n### Professional Worksheet Generator — Grades 1 to 10")
    gr.Markdown(AI_STATUS_BANNER)

    with gr.Accordion("🧭 Getting Started — New user? Click here / Naya user yahan click karein", open=True):
        gr.Markdown(GETTING_STARTED_GUIDE)

    questions_state = gr.State([])  # holds the last generated List[QuestionItem]

    # ---------------- Branding ----------------
    with gr.Accordion("🏫 1. School Branding & Cover Page", open=True):
        with gr.Row():
            school_name = gr.Textbox(label="School Name", placeholder="e.g. Roots International Schools")
            logo_file = gr.File(label="Upload School Logo", file_types=["image"])
            logo_size = gr.Slider(20, 150, value=60, step=5, label="Logo Size (pt)")
        cover_color = gr.ColorPicker(label="Cover / Header Theme Color", value="#2563EB")

    # ---------------- Shared Controls ----------------
    with gr.Accordion("⚙️ 2. Worksheet Settings (apply to both methods)", open=True):
        with gr.Row():
            grade = gr.Dropdown(GRADES, value="3", label="Grade")
            subject = gr.Dropdown(SUBJECTS, value="English", label="Subject", allow_custom_value=True)
            topic = gr.Textbox(label="Topic", placeholder="e.g. Nouns, Photosynthesis, Fractions")
            language = gr.Radio(["English", "Urdu"], value="English", label="Worksheet Language")
        with gr.Row():
            q_types = gr.CheckboxGroup(QUESTION_TYPES, value=["Choose the Best Option (MCQ)"], label="Question Types (select any combination)")
            num_q = gr.Number(value=10, precision=0, label="Total Number of Questions")
        with gr.Row():
            q_per_page = gr.Number(value=0, precision=0, label="Questions per Page (0 = unlimited/auto-flow)")
            spacing = gr.Slider(0, 60, value=18, step=2, label="Spacing Between Questions (pt)")
            binding_gap = gr.Slider(0, 40, value=0, step=2, label="Binding Gap (mm, extra left margin)")
        with gr.Row():
            num_copies = gr.Number(value=1, precision=0, label="Number of Copies")
            copy_mode = gr.Radio(["Identical", "Randomized"], value="Identical", label="Copy Mode")
            visual_theme = gr.Radio(THEMES, value="Auto (based on grade)", label="Visual Theme")
        with gr.Row():
            show_name = gr.Checkbox(value=True, label="Show Student Name field")
            show_class = gr.Checkbox(value=True, label="Show Class field")
            show_roll = gr.Checkbox(value=True, label="Show Roll No. field")
            show_date = gr.Checkbox(value=True, label="Show Date field")

    shared_inputs = [school_name, logo_file, logo_size, cover_color,
                      grade, subject, topic, language,
                      q_types, num_q,
                      q_per_page, spacing, binding_gap, num_copies, copy_mode,
                      visual_theme,
                      show_name, show_class, show_roll, show_date]

    # ---------------- Method 1 ----------------
    with gr.Accordion("📂 3. Method 1 — Curriculum-Based Generation", open=False):
        gr.Markdown("Use **either** Path A or Path B. Use the 'Clear' button to reset a path you don't need.")
        with gr.Tab("Path A — Upload File + Chapter/Topic/Page"):
            m1a_file = gr.File(label="Upload Curriculum File (PDF/DOCX/TXT)")
            with gr.Row():
                m1a_chapter = gr.Textbox(label="Chapter")
                m1a_topic = gr.Textbox(label="Topic")
                m1a_page = gr.Textbox(label="Page No.")
            with gr.Row():
                m1a_generate = gr.Button("Generate from File ✨", variant="primary")
                m1a_clear = gr.Button("Clear Path A")
            m1a_preview = gr.Markdown()

        with gr.Tab("Path B — Grade / Subject / Topic Dropdown"):
            gr.Markdown("Uses the Grade / Subject / Topic selected above in Worksheet Settings.")
            with gr.Row():
                m1b_generate = gr.Button("Generate from Selection ✨", variant="primary")
                m1b_clear = gr.Button("Clear Path B")
            m1b_preview = gr.Markdown()

    # ---------------- Method 2 ----------------
    with gr.Accordion("📝 4. Method 2 — Paste or Upload Any Content", open=False):
        m2_text = gr.Textbox(label="Paste content here (e.g. from ChatGPT, notes, textbook excerpt)", lines=8)
        m2_file = gr.File(label="Or upload a file instead")
        with gr.Row():
            m2_generate = gr.Button("Clean & Generate ✨", variant="primary")
            m2_clear = gr.Button("Clear")
        m2_preview = gr.Markdown()

    # ---------------- Method 3 ----------------
    with gr.Accordion("🗂️ 5. Method 3 — Your Own Question Bank (no API key / no cost)", open=(not AI_ENABLED)):
        gr.Markdown(
            "Upload a CSV or Excel file with columns: `grade, subject, topic, q_type, text, options, "
            "match_left, match_right, answer, image_keyword`. Use `|` to separate multiple options "
            "(e.g. `Cat|Dog|Cow|Horse`). This works completely offline from AI — matches are filtered "
            "using the Grade / Subject / Topic / Question Types set above in Worksheet Settings."
        )
        m3_file = gr.File(label="Upload Question Bank (CSV or XLSX)")
        with gr.Row():
            m3_generate = gr.Button("Load & Build Worksheet ✨", variant="primary")
            m3_clear = gr.Button("Clear")
        m3_preview = gr.Markdown()

    # ---------------- Preview & Export ----------------
    with gr.Accordion("👁️ 6. Preview & Export", open=True):
        export_format = gr.Radio(["PDF", "Word"], value="PDF", label="Export Format")
        create_btn = gr.Button("📄 Create Worksheet", variant="primary", size="lg")
        status = gr.Markdown()
        download_file = gr.File(label="Download Your Worksheet")

    # ---------------- Wiring ----------------
    m1a_generate.click(generate_method1_pathA,
                        inputs=[m1a_file, m1a_chapter, m1a_topic, m1a_page] + shared_inputs,
                        outputs=[m1a_preview, questions_state])
    m1a_clear.click(lambda: (None, "", "", "", "", []), outputs=[m1a_file, m1a_chapter, m1a_topic, m1a_page, m1a_preview, questions_state])

    m1b_generate.click(generate_method1_pathB, inputs=shared_inputs, outputs=[m1b_preview, questions_state])
    m1b_clear.click(lambda: ("", []), outputs=[m1b_preview, questions_state])

    m2_generate.click(generate_method2, inputs=[m2_text, m2_file] + shared_inputs, outputs=[m2_preview, questions_state])
    m2_clear.click(lambda: ("", None, "", []), outputs=[m2_text, m2_file, m2_preview, questions_state])

    m3_generate.click(generate_method3, inputs=[m3_file] + shared_inputs, outputs=[m3_preview, questions_state])
    m3_clear.click(lambda: (None, "", []), outputs=[m3_file, m3_preview, questions_state])

    create_btn.click(export_worksheet, inputs=[questions_state, export_format] + shared_inputs,
                      outputs=[download_file, status])

app.launch()
