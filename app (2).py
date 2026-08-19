"""
SchoolSheet Studio — Professional Worksheet Generator (Streamlit version)
Deployable for free on Streamlit Community Cloud (share.streamlit.io) —
no credit card, connects directly to a GitHub repo.
"""
import os
import requests

# ------------------------------------------------------------------
# AI Provider & Keys — configured via Streamlit "Secrets"
# (App settings -> Secrets, in TOML format, e.g.:
#    ANTHROPIC_API_KEY = "sk-ant-..."
#    AI_PROVIDER = "anthropic"
#  Streamlit loads secrets into os.environ automatically at startup
#  when they are also mirrored here, so we read from both.)
# ------------------------------------------------------------------
try:
    import streamlit as st
    for _k in ("ANTHROPIC_API_KEY", "GEMINI_API_KEY", "PIXABAY_API_KEY", "AI_PROVIDER"):
        if _k in st.secrets and not os.environ.get(_k):
            os.environ[_k] = str(st.secrets[_k])
except Exception:
    pass

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "").strip()
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
PIXABAY_API_KEY = os.environ.get("PIXABAY_API_KEY", "").strip()
_provider_pref = os.environ.get("AI_PROVIDER", "").strip().lower()

if _provider_pref == "anthropic" and ANTHROPIC_API_KEY:
    AI_ENABLED = True
    ACTIVE_PROVIDER = "anthropic"
elif _provider_pref == "gemini" and GEMINI_API_KEY:
    AI_ENABLED = True
    ACTIVE_PROVIDER = "gemini"
elif _provider_pref in ("none", ""):
    if ANTHROPIC_API_KEY:
        AI_ENABLED = True
        ACTIVE_PROVIDER = "anthropic"
    elif GEMINI_API_KEY:
        AI_ENABLED = True
        ACTIVE_PROVIDER = "gemini"
    else:
        AI_ENABLED = False
        ACTIVE_PROVIDER = None
else:
    AI_ENABLED = False
    ACTIVE_PROVIDER = None

# ------------------------------------------------------------------
# Urdu Unicode font (Noto Nastaliq Urdu) so Urdu text renders correctly
# in PDF and Word output.
# ------------------------------------------------------------------
FONT_DIR = "./fonts"
os.makedirs(FONT_DIR, exist_ok=True)
URDU_FONT_PATH = os.path.join(FONT_DIR, "NotoNastaliqUrdu-Regular.ttf")

if not os.path.exists(URDU_FONT_PATH):
    font_url = "https://raw.githubusercontent.com/google/fonts/main/ofl/notonastaliqurdu/NotoNastaliqUrdu%5Bwght%5D.ttf"
    try:
        r = requests.get(font_url, timeout=30)
        r.raise_for_status()
        with open(URDU_FONT_PATH, "wb") as f:
            f.write(r.content)
    except Exception:
        URDU_FONT_PATH = None

LATIN_FONT_PATH = None
try:
    import matplotlib
    LATIN_FONT_PATH = os.path.join(os.path.dirname(matplotlib.__file__), "mpl-data", "fonts", "ttf", "DejaVuSans.ttf")
    if not os.path.exists(LATIN_FONT_PATH):
        LATIN_FONT_PATH = None
except Exception:
    LATIN_FONT_PATH = None

import pdfplumber
import docx as docx_lib

# ------------------------------------------------------------------
# Unified AI call layer — one function, works with EITHER provider.
# Every other step in this notebook calls _call_ai() and never talks
# to the Anthropic or Gemini SDKs directly, so switching providers in
# Step 2 doesn't require touching any code below.
# ------------------------------------------------------------------
_anthropic_client = None
_gemini_model = None

CLAUDE_MODEL = "claude-sonnet-4-6"
GEMINI_MODEL = "gemini-2.5-flash"

if AI_ENABLED and ACTIVE_PROVIDER == "anthropic":
    from anthropic import Anthropic
    _anthropic_client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
elif AI_ENABLED and ACTIVE_PROVIDER == "gemini":
    import google.generativeai as genai
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    _gemini_model = genai.GenerativeModel(GEMINI_MODEL)


def _call_ai(system_prompt: str, user_content: str, max_tokens: int = 4000) -> str:
    """
    Single entry point for every AI call in this notebook.
    Sends `user_content` with `system_prompt` as instructions and returns
    the raw text reply — regardless of which provider is active.
    Raises RuntimeError if AI is not enabled (no valid key/provider).
    """
    if not AI_ENABLED:
        raise RuntimeError(
            "This feature isn't available right now — please use 'Method 3: Your Own Question Bank' instead."
        )
    if ACTIVE_PROVIDER == "anthropic":
        resp = _anthropic_client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=max_tokens,
            system=system_prompt,
            messages=[{"role": "user", "content": user_content}],
        )
        return "".join(b.text for b in resp.content if b.type == "text").strip()
    elif ACTIVE_PROVIDER == "gemini":
        full_prompt = f"{system_prompt}\n\n{user_content}"
        resp = _gemini_model.generate_content(
            full_prompt,
            generation_config={"max_output_tokens": max_tokens},
        )
        return (getattr(resp, "text", None) or "").strip()
    else:
        raise RuntimeError(f"Unknown AI provider: {ACTIVE_PROVIDER}")


def extract_text_from_file(file_path: str) -> str:
    """Reads .pdf, .docx or .txt and returns raw text."""
    if file_path is None:
        return ""
    lower = file_path.lower()
    try:
        if lower.endswith(".pdf"):
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    text_parts.append(t)
            return "\n".join(text_parts)
        elif lower.endswith(".docx"):
            d = docx_lib.Document(file_path)
            return "\n".join(p.text for p in d.paragraphs)
        else:  # .txt or anything else — read as plain text
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        return f"[ERROR READING FILE: {e}]"


def clean_pasted_content(raw_text: str) -> str:
    """
    Removes AI chatter / preambles / postambles / instructions that are
    NOT actual question content (e.g. "Here are your questions:",
    "I hope this helps!", numbering artifacts from chat exports, etc.)
    Keeps only the real educational content, unaltered.
    """
    if not raw_text or not raw_text.strip():
        return ""

    if not AI_ENABLED:
        # Simple rule-based fallback cleaning when no API key is configured:
        # drops common AI-chatter lines and blank/short filler lines.
        junk_markers = [
            "here are your questions", "here's your worksheet", "i hope this helps",
            "hope this helps", "let me know if", "feel free to", "sure, here",
            "certainly!", "as requested", "best regards", "regards,",
        ]
        lines = raw_text.splitlines()
        kept = []
        for line in lines:
            l = line.strip()
            if not l:
                continue
            if any(marker in l.lower() for marker in junk_markers):
                continue
            kept.append(line)
        return "\n".join(kept).strip()

    system_prompt = (
        "You are a strict text-cleaning engine for a school worksheet system. "
        "You will be given raw text that may contain educational questions or "
        "content mixed with irrelevant chatter (greetings, meta-comments like "
        "'Here are your questions', 'I hope this helps', signatures, extra "
        "headers, disclaimers, or duplicate/broken lines). "
        "Return ONLY the genuine educational content — questions, passages, "
        "topics, instructions meant for students — exactly as originally "
        "written, with no rephrasing. Remove every non-content line. "
        "Do not add any commentary, notes, or explanation of your own. "
        "Output nothing except the cleaned content."
    )

    return _call_ai(system_prompt, raw_text, max_tokens=4000)


print("Text extraction & cleaning utilities ready. Active AI provider:", ACTIVE_PROVIDER or "None (Question Bank mode)")

from dataclasses import dataclass, field
from typing import List, Optional
import json, random

QUESTION_TYPES = [
    "Choose the Best Option (MCQ)",
    "Fill in the Blanks",
    "Match the Column",
    "Short Question",
    "Detailed Question",
]

GRADES = [str(g) for g in range(1, 11)]  # Grade 1 to Grade 10

THEMES = ["Auto (based on grade)", "Colorful / Playful", "Professional / Formal"]


@dataclass
class QuestionItem:
    q_type: str
    text: str
    options: Optional[List[str]] = None       # for MCQ
    match_left: Optional[List[str]] = None     # for Match the Column
    match_right: Optional[List[str]] = None
    answer: Optional[str] = None
    image_keyword: Optional[str] = None        # e.g. "apple", "cat" for young grades

    def to_dict(self):
        return self.__dict__


@dataclass
class WorksheetConfig:
    school_name: str = "Your School Name"
    logo_path: Optional[str] = None
    logo_size_pt: int = 60          # logo width in points
    cover_color_hex: str = "#2563EB"
    font_theme: str = "Poppins/Default"

    grade: str = "3"
    subject: str = "General"
    topic: str = "General Topic"
    language: str = "English"       # "English" or "Urdu"

    question_types: List[str] = field(default_factory=lambda: ["Choose the Best Option (MCQ)"])
    num_questions: int = 10

    questions_per_page: Optional[int] = None   # None = unlimited/auto-flow
    spacing_pt: int = 18                        # vertical space between questions
    binding_gap_mm: int = 0                      # extra margin for binding
    num_copies: int = 1
    copy_mode: str = "Identical"                 # "Identical" or "Randomized"

    visual_theme: str = "Auto (based on grade)"

    show_student_name: bool = True
    show_class: bool = True
    show_roll_no: bool = True
    show_date: bool = True
    include_answer_key: bool = True

    def effective_theme(self) -> str:
        if self.visual_theme != "Auto (based on grade)":
            return self.visual_theme
        return "Colorful / Playful" if int(self.grade) <= 4 else "Professional / Formal"


print("Data models ready. Question types:", QUESTION_TYPES)

import pandas as pd

def _build_generation_prompt(cfg: "WorksheetConfig", source_text: Optional[str]) -> str:
    types_str = ", ".join(cfg.question_types)
    lang_line = (
        "Write the ENTIRE worksheet content in proper, correct, natural Urdu "
        "(Urdu script, not Roman Urdu)."
        if cfg.language == "Urdu"
        else "Write everything in clear, simple, grade-appropriate English."
    )
    source_block = (
        f"Base every question strictly on the following source content "
        f"(do not invent facts outside it):\n\n{source_text}\n"
        if source_text
        else f"Topic: {cfg.topic}\nSubject: {cfg.subject}\n"
    )
    image_line = (
        "Since this is for a young grade (1-4), for at least half the "
        "questions add a simple, common, real-world 'image_keyword' "
        "(e.g. apple, cat, monkey, ball, sun, book) that a teacher could "
        "illustrate the question with. Use only simple everyday objects/animals."
        if int(cfg.grade) <= 4
        else "This is for an older grade — do not add image_keyword, keep it null."
    )

    return f"""
You are an expert school curriculum question-writer for Grade {cfg.grade}, subject {cfg.subject}.
{source_block}
Generate exactly {cfg.num_questions} TOTAL questions, using ONLY these question types: {types_str}.
Distribute the questions across the selected types as evenly as possible.
{lang_line}
{image_line}

Rules:
- Every question must be completely different in content from every other question (no repeats, no rephrasing of the same fact twice).
- For "Choose the Best Option (MCQ)": provide exactly 4 options and mark the correct answer.
- For "Fill in the Blanks": use a blank shown as "_____" inside the sentence, and give the answer.
- For "Match the Column": provide a left list and right list (same length, 4-6 pairs), shuffle the right list order, and give the answer as pairs.
- For "Short Question": a question needing a 1-3 sentence answer.
- For "Detailed Question": a question needing a paragraph-length answer.
- Keep difficulty appropriate for Grade {cfg.grade}.

Return ONLY valid JSON (no markdown fences, no commentary), as a list of objects with this exact schema:
[
  {{
    "q_type": "one of: Choose the Best Option (MCQ) | Fill in the Blanks | Match the Column | Short Question | Detailed Question",
    "text": "the question text (or sentence with blank)",
    "options": ["...", "...", "...", "..."] or null,
    "match_left": ["..."] or null,
    "match_right": ["..."] or null,
    "answer": "the correct answer / answer key",
    "image_keyword": "simple keyword" or null
  }}
]
""".strip()


def generate_questions_ai(cfg: "WorksheetConfig", source_text: Optional[str] = None) -> List[QuestionItem]:
    """Calls the active AI provider (Anthropic or Gemini) to generate a fresh, non-repeating question set."""
    if not AI_ENABLED:
        raise RuntimeError(
            "This feature isn't available right now — please use 'Method 3: Your Own Question Bank' instead."
        )
    prompt = _build_generation_prompt(cfg, source_text)
    raw = _call_ai(
        system_prompt=(
            "You are an expert school curriculum question-writer. "
            "Return ONLY valid JSON (no markdown fences, no commentary)."
        ),
        user_content=prompt,
        max_tokens=4000,
    )
    raw = raw.strip("`")
    if raw.lower().startswith("json"):
        raw = raw[4:].strip()
    try:
        data = json.loads(raw)
    except Exception:
        # last-resort: try to locate the JSON array inside the text
        start = raw.find("[")
        end = raw.rfind("]")
        data = json.loads(raw[start:end + 1])

    return [QuestionItem(
        q_type=item.get("q_type", "Short Question"),
        text=item.get("text", ""),
        options=item.get("options"),
        match_left=item.get("match_left"),
        match_right=item.get("match_right"),
        answer=item.get("answer"),
        image_keyword=item.get("image_keyword"),
    ) for item in data]


def load_question_bank(file_path: str) -> pd.DataFrame:
    """
    Optional plug-in question bank. Expected columns (csv or xlsx):
    grade, subject, topic, q_type, text, options, match_left, match_right, answer, image_keyword
    'options' / 'match_left' / 'match_right' cells should be pipe-separated, e.g.  "A|B|C|D"
    """
    if file_path.lower().endswith(".csv"):
        return pd.read_csv(file_path)
    return pd.read_excel(file_path)


def questions_from_bank(df: "pd.DataFrame", cfg: "WorksheetConfig") -> List[QuestionItem]:
    filtered = df[
        (df["grade"].astype(str) == str(cfg.grade))
        & (df["subject"].str.lower() == cfg.subject.lower())
        & (df["topic"].str.lower() == cfg.topic.lower())
        & (df["q_type"].isin(cfg.question_types))
    ]
    filtered = filtered.sample(frac=1).head(cfg.num_questions)  # shuffle + limit
    items = []
    for _, row in filtered.iterrows():
        split = lambda v: str(v).split("|") if pd.notna(v) else None
        items.append(QuestionItem(
            q_type=row["q_type"],
            text=row["text"],
            options=split(row.get("options")),
            match_left=split(row.get("match_left")),
            match_right=split(row.get("match_right")),
            answer=row.get("answer"),
            image_keyword=row.get("image_keyword") if pd.notna(row.get("image_keyword")) else None,
        ))
    return items


def randomize_copy(questions: List[QuestionItem]) -> List[QuestionItem]:
    """Produces a shuffled variant (question order + MCQ option order) for randomized copies."""
    import copy as _copy
    shuffled = _copy.deepcopy(questions)
    random.shuffle(shuffled)
    for q in shuffled:
        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            correct = q.options[0] if q.answer == q.options[0] else q.answer
            opts = q.options[:]
            random.shuffle(opts)
            q.options = opts
    return shuffled


print("AI generation engine + question bank plug-in ready.")

import hashlib

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_OK = True
except Exception as e:
    print("⚠️ Pillow (PIL) failed to import:", e)
    print("   Images will be skipped, but the rest of the app (PDF/Word generation,")
    print("   AI question generation, all worksheet features) will still work fine.")
    print("   To fix images: Runtime -> Restart session, then run all cells again from Step 1.")
    PIL_OK = False

IMAGE_CACHE_DIR = "./wg_images"
os.makedirs(IMAGE_CACHE_DIR, exist_ok=True)

_FALLBACK_COLORS = [
    "#FF6B6B", "#4ECDC4", "#FFD93D", "#6BCB77", "#4D96FF", "#FF922B", "#C084FC"
]


def _cache_path(keyword: str) -> str:
    h = hashlib.md5(keyword.encode()).hexdigest()[:10]
    return os.path.join(IMAGE_CACHE_DIR, f"{h}.png")


def _make_fallback_icon(keyword: str, path: str, size: int = 300):
    """Generates a clean, simple flat-color icon with the keyword label,
    used only if no Pixabay key / no internet result is available."""
    if not PIL_OK:
        return None
    color = _FALLBACK_COLORS[hash(keyword) % len(_FALLBACK_COLORS)]
    img = Image.new("RGB", (size, size), "white")
    draw = ImageDraw.Draw(img)
    draw.ellipse([20, 20, size - 20, size - 20], fill=color)
    try:
        font = ImageFont.truetype(LATIN_FONT_PATH, 28) if LATIN_FONT_PATH else ImageFont.load_default()
    except Exception:
        font = ImageFont.load_default()
    text = keyword.upper()
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((size - tw) / 2, size - 55), text, fill="white", font=font)
    img.save(path)
    return path


def get_image_for_keyword(keyword: str) -> Optional[str]:
    """Returns a local file path to an image representing `keyword`.
    Tries Pixabay (real photo) first if a key is set, else falls back
    to a clean generated icon so the worksheet is never left blank.
    Returns None (image simply skipped) if Pillow is unavailable."""
    if not keyword:
        return None
    path = _cache_path(keyword)
    if os.path.exists(path):
        return path

    pixabay_key = os.environ.get("PIXABAY_API_KEY", "").strip()
    if pixabay_key:
        try:
            r = requests.get(
                "https://pixabay.com/api/",
                params={
                    "key": pixabay_key,
                    "q": keyword,
                    "image_type": "photo",
                    "category": "animals,food,education",
                    "safesearch": "true",
                    "per_page": 3,
                },
                timeout=15,
            )
            data = r.json()
            if data.get("hits"):
                img_url = data["hits"][0]["webformatURL"]
                img_data = requests.get(img_url, timeout=15).content
                with open(path, "wb") as f:
                    f.write(img_data)
                return path
        except Exception as e:
            print(f"Pixabay fetch failed for '{keyword}', using fallback icon. ({e})")

    return _make_fallback_icon(keyword, path)


print("Image engine ready (real photos via Pixabay, with automatic icon fallback).")

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display

# Register fonts once
_FONTS_REGISTERED = False

def _register_fonts():
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    if URDU_FONT_PATH and os.path.exists(URDU_FONT_PATH):
        pdfmetrics.registerFont(TTFont("Urdu", URDU_FONT_PATH))
    if LATIN_FONT_PATH and os.path.exists(LATIN_FONT_PATH):
        pdfmetrics.registerFont(TTFont("Latin", LATIN_FONT_PATH))
    _FONTS_REGISTERED = True

def _shape(text: str, is_urdu: bool) -> str:
    if not is_urdu:
        return text
    reshaped = arabic_reshaper.reshape(text)
    return get_display(reshaped)


THEME_COLORS = {
    "Colorful / Playful": {"primary": "#FF6B6B", "secondary": "#FFD93D", "accent": "#4ECDC4"},
    "Professional / Formal": {"primary": "#1E3A8A", "secondary": "#334155", "accent": "#2563EB"},
}


def render_pdf(cfg: "WorksheetConfig", questions: List[QuestionItem], output_path: str):
    _register_fonts()
    is_urdu = cfg.language == "Urdu"
    body_font = "Urdu" if (is_urdu and "Urdu" in pdfmetrics.getRegisteredFontNames()) else "Helvetica"
    head_font = "Latin" if "Latin" in pdfmetrics.getRegisteredFontNames() else "Helvetica-Bold"

    theme = THEME_COLORS[cfg.effective_theme()]
    page_w, page_h = A4
    left_margin = 18 * mm + cfg.binding_gap_mm * mm
    right_margin = 18 * mm
    top_margin = 15 * mm
    bottom_margin = 15 * mm

    c = canvas.Canvas(output_path, pagesize=A4)

    def draw_header(page_no_label=""):
        y = page_h - top_margin
        # Logo
        if cfg.logo_path and os.path.exists(cfg.logo_path):
            try:
                c.drawImage(cfg.logo_path, left_margin, y - cfg.logo_size_pt,
                            width=cfg.logo_size_pt, height=cfg.logo_size_pt,
                            preserveAspectRatio=True, mask="auto")
            except Exception:
                pass
        # School name
        c.setFont(head_font, 16)
        c.setFillColor(theme["primary"])
        c.drawCentredString(page_w / 2, y - 20, _shape(cfg.school_name, False))
        c.setFont(head_font, 11)
        c.setFillColor("#333333")
        c.drawCentredString(page_w / 2, y - 38, f"Grade {cfg.grade}  |  {cfg.subject}  |  {cfg.topic}")

        # divider line
        c.setStrokeColor(theme["accent"])
        c.setLineWidth(1.5)
        c.line(left_margin, y - 48, page_w - right_margin, y - 48)

        # student info fields
        fy = y - 65
        c.setFont("Helvetica", 10)
        c.setFillColor("#000000")
        fields = []
        if cfg.show_student_name: fields.append("Student Name: ______________________")
        if cfg.show_class: fields.append("Class: _______")
        if cfg.show_roll_no: fields.append("Roll No: _______")
        if cfg.show_date: fields.append("Date: _______")
        line1 = "     ".join(fields[:2])
        line2 = "     ".join(fields[2:])
        c.drawString(left_margin, fy, line1)
        if line2:
            c.drawString(left_margin, fy - 14, line2)

        c.setStrokeColor("#999999")
        c.setLineWidth(0.7)
        c.line(left_margin, fy - 24, page_w - right_margin, fy - 24)
        return fy - 40  # y position where question body should start

    def new_page():
        c.showPage()
        return draw_header()

    y_cursor = draw_header()
    q_count_on_page = 0
    max_width = page_w - left_margin - right_margin

    def wrap_text(text, font, size, max_w):
        c.setFont(font, size)
        words = text.split(" ")
        lines, cur = [], ""
        for w in words:
            trial = (cur + " " + w).strip()
            if c.stringWidth(trial, font, size) <= max_w:
                cur = trial
            else:
                lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
        return lines

    for idx, q in enumerate(questions, start=1):
        needed_space = 70  # rough estimate; grows below per type
        if y_cursor < bottom_margin + needed_space or (
            cfg.questions_per_page and q_count_on_page >= cfg.questions_per_page
        ):
            y_cursor = new_page()
            q_count_on_page = 0

        c.setFont(body_font, 11)
        c.setFillColor("#000000")
        prefix = f"{idx}. "
        q_text = _shape(prefix + q.text, is_urdu)
        lines = wrap_text(q_text, body_font, 11, max_width)
        for line in lines:
            c.drawString(left_margin, y_cursor, line)
            y_cursor -= 15

        # image for young grades
        if q.image_keyword:
            img_path = get_image_for_keyword(q.image_keyword)
            if img_path:
                try:
                    c.drawImage(img_path, left_margin, y_cursor - 60, width=55, height=55,
                                preserveAspectRatio=True, mask="auto")
                    y_cursor -= 65
                except Exception:
                    pass

        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            for i, opt in enumerate(q.options):
                letter = chr(97 + i)
                c.drawString(left_margin + 15, y_cursor, _shape(f"({letter}) {opt}", is_urdu))
                y_cursor -= 14

        elif q.q_type == "Match the Column" and q.match_left and q.match_right:
            col2_x = left_margin + max_width / 2
            rows = max(len(q.match_left), len(q.match_right))
            for i in range(rows):
                left_txt = q.match_left[i] if i < len(q.match_left) else ""
                right_txt = q.match_right[i] if i < len(q.match_right) else ""
                c.drawString(left_margin + 15, y_cursor, _shape(f"{i+1}. {left_txt}", is_urdu))
                c.drawString(col2_x, y_cursor, _shape(f"{chr(97+i)}. {right_txt}", is_urdu))
                y_cursor -= 14

        elif q.q_type == "Short Question":
            for _ in range(2):
                y_cursor -= 14
                c.setStrokeColor("#cccccc")
                c.line(left_margin, y_cursor, page_w - right_margin, y_cursor)

        elif q.q_type == "Detailed Question":
            for _ in range(5):
                y_cursor -= 14
                c.setStrokeColor("#cccccc")
                c.line(left_margin, y_cursor, page_w - right_margin, y_cursor)

        elif q.q_type == "Fill in the Blanks":
            pass  # blank already embedded in text

        y_cursor -= cfg.spacing_pt
        q_count_on_page += 1

    # ---- Answer Key (separate page at the end) ----
    if getattr(cfg, "include_answer_key", True):
        c.showPage()
        y_cursor = draw_header()
        c.setFont(head_font, 14)
        c.setFillColor(theme["primary"])
        c.drawCentredString(page_w / 2, y_cursor + 10, "Answer Key")
        y_cursor -= 15
        c.setFont(body_font, 11)
        c.setFillColor("#000000")
        for idx, q in enumerate(questions, start=1):
            if y_cursor < bottom_margin + 20:
                y_cursor = new_page()
            ans_text = q.answer if q.answer else "(open-ended — no fixed answer)"
            line = _shape(f"{idx}. {ans_text}", is_urdu)
            for wrapped in wrap_text(line, body_font, 11, max_width):
                c.drawString(left_margin, y_cursor, wrapped)
                y_cursor -= 15

    c.save()
    return output_path


def build_pdf_copies(cfg: "WorksheetConfig", base_questions: List[QuestionItem], output_dir: str) -> List[str]:
    os.makedirs(output_dir, exist_ok=True)
    paths = []
    for i in range(1, cfg.num_copies + 1):
        qs = randomize_copy(base_questions) if cfg.copy_mode == "Randomized" else base_questions
        out = os.path.join(output_dir, f"worksheet_copy_{i}.pdf")
        render_pdf(cfg, qs, out)
        paths.append(out)
    return paths


print("PDF renderer ready.")

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _apply_urdu_font(run, size=12):
    run.font.name = "Noto Nastaliq Urdu"
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), "Noto Nastaliq Urdu")


def render_docx(cfg: "WorksheetConfig", questions: List[QuestionItem], output_path: str):
    is_urdu = cfg.language == "Urdu"
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Mm(18 + cfg.binding_gap_mm)
    section.right_margin = Mm(18)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    theme_primary = RGBColor(0x1E, 0x3A, 0x8A) if cfg.effective_theme() == "Professional / Formal" else RGBColor(0xE0, 0x50, 0x50)

    # Logo + School name
    if cfg.logo_path and os.path.exists(cfg.logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(cfg.logo_path, width=Pt(cfg.logo_size_pt))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(cfg.school_name)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = theme_primary

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Grade {cfg.grade}  |  {cfg.subject}  |  {cfg.topic}").font.size = Pt(11)

    doc.add_paragraph("_" * 90)

    # Student fields
    fields = []
    if cfg.show_student_name: fields.append("Student Name: ______________________")
    if cfg.show_class: fields.append("Class: _______")
    if cfg.show_roll_no: fields.append("Roll No: _______")
    if cfg.show_date: fields.append("Date: _______")
    if fields:
        doc.add_paragraph("     ".join(fields[:2]))
        if len(fields) > 2:
            doc.add_paragraph("     ".join(fields[2:]))
    doc.add_paragraph("_" * 90)

    for idx, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        if is_urdu:
            _set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{idx}. {q.text}")
        run.bold = True
        if is_urdu:
            _apply_urdu_font(run, 13)
        else:
            run.font.size = Pt(12)

        if q.image_keyword:
            img_path = get_image_for_keyword(q.image_keyword)
            if img_path:
                try:
                    doc.add_picture(img_path, width=Inches(0.7))
                except Exception:
                    pass

        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            for i, opt in enumerate(q.options):
                op = doc.add_paragraph(f"    ({chr(97+i)}) {opt}")
                if is_urdu:
                    _set_rtl(op)
                    op.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _apply_urdu_font(op.runs[0], 12)

        elif q.q_type == "Match the Column" and q.match_left and q.match_right:
            table = doc.add_table(rows=max(len(q.match_left), len(q.match_right)), cols=2)
            table.style = "Table Grid"
            for i in range(len(table.rows)):
                l = q.match_left[i] if i < len(q.match_left) else ""
                r_ = q.match_right[i] if i < len(q.match_right) else ""
                table.cell(i, 0).text = f"{i+1}. {l}"
                table.cell(i, 1).text = f"{chr(97+i)}. {r_}"

        elif q.q_type == "Short Question":
            for _ in range(2):
                doc.add_paragraph("_" * 90)

        elif q.q_type == "Detailed Question":
            for _ in range(5):
                doc.add_paragraph("_" * 90)

        # spacing between questions
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(cfg.spacing_pt)

        if cfg.questions_per_page and idx % cfg.questions_per_page == 0 and idx != len(questions):
            doc.add_page_break()

    if getattr(cfg, "include_answer_key", True):
        doc.add_page_break()
        ak_title = doc.add_paragraph()
        ak_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ak_run = ak_title.add_run("Answer Key")
        ak_run.bold = True
        ak_run.font.size = Pt(16)
        ak_run.font.color.rgb = theme_primary

        for idx, q in enumerate(questions, start=1):
            ans_text = q.answer if q.answer else "(open-ended — no fixed answer)"
            p = doc.add_paragraph()
            if is_urdu:
                _set_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"{idx}. {ans_text}")
            if is_urdu:
                _apply_urdu_font(run, 12)
            else:
                run.font.size = Pt(11)

    doc.save(output_path)
    return output_path


def build_docx_copies(cfg: "WorksheetConfig", base_questions: List[QuestionItem], output_dir: str) -> List[str]:
    os.makedirs(output_dir, exist_ok=True)
    paths = []
    for i in range(1, cfg.num_copies + 1):
        qs = randomize_copy(base_questions) if cfg.copy_mode == "Randomized" else base_questions
        out = os.path.join(output_dir, f"worksheet_copy_{i}.docx")
        render_docx(cfg, qs, out)
        paths.append(out)
    return paths


print("DOCX renderer ready.")

# =====================================================================================
#  STREAMLIT UI
# =====================================================================================
import streamlit as st
import tempfile, shutil

APP_NAME = "SchoolSheet Studio"
WORK_DIR = "./wg_output"
os.makedirs(WORK_DIR, exist_ok=True)

SUBJECTS = ["English", "Urdu", "Mathematics", "Science", "Social Studies", "Islamiyat", "Computer Science", "General Knowledge"]

st.set_page_config(page_title=APP_NAME, page_icon="\U0001F4D8", layout="wide")

# =====================================================================================
#  THEME — white + navy blue, premium/professional look
# =====================================================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;500;600;700;800&display=swap');

html, body, [class*="css"]  { font-family: 'Poppins', sans-serif; }

/* Main app background — clean white/off-white */
.stApp {
    background: #F5F7FB;
}

/* Sidebar — deep navy blue */
section[data-testid="stSidebar"] {
    background: #0B1E3F;
    border-right: 1px solid rgba(255,255,255,0.06);
}
section[data-testid="stSidebar"] * { color: #E7ECF7 !important; }
section[data-testid="stSidebar"] label { color: #E7ECF7 !important; }

/* Headings / body text on the light canvas */
h1, h2, h3, h4, h5, h6 { color: #0B1E3F !important; font-weight: 700 !important; }
p, span, label, .stMarkdown, .stCaption { color: #3A4258; }

/* Card container helper — wrap sections in st.container(border=True) to get this look */
div[data-testid="stVerticalBlockBorderWrapper"] {
    background: #FFFFFF;
    border-radius: 16px;
    padding: 4px;
    border: 1px solid #E6E9F2 !important;
    box-shadow: 0 6px 20px rgba(11,30,63,0.08);
}
div[data-testid="stVerticalBlockBorderWrapper"] * {
    color: #1B2A4A !important;
}
div[data-testid="stVerticalBlockBorderWrapper"] h1,
div[data-testid="stVerticalBlockBorderWrapper"] h2,
div[data-testid="stVerticalBlockBorderWrapper"] h3 { color: #0B1E3F !important; }

/* Buttons */
.stButton > button {
    background: #14356E;
    color: #FFFFFF;
    border-radius: 10px;
    border: none;
    font-weight: 600;
    padding: 0.5em 1.2em;
}
.stButton > button:hover { background: #0B1E3F; color: #FFFFFF; }
.stButton > button[kind="primary"] { background: #14356E; }
.stButton > button p, .stButton > button div, .stButton > button span {
    color: #FFFFFF !important;
}

/* Inputs */
.stTextInput input, .stSelectbox div[data-baseweb="select"], .stNumberInput input, .stTextArea textarea {
    border-radius: 8px !important;
    border: 1px solid #D7DCE8 !important;
}

/* Metric cards */
div[data-testid="stMetric"] {
    background: #FFFFFF;
    border-radius: 14px;
    padding: 14px 16px;
    border: 1px solid #E6E9F2;
    box-shadow: 0 4px 14px rgba(11,30,63,0.06);
}
div[data-testid="stMetric"] label, div[data-testid="stMetric"] div { color: #0B1E3F !important; }

/* Tabs */
.stTabs [data-baseweb="tab"] { color: #3A4258; font-weight: 600; }
.stTabs [aria-selected="true"] { color: #0B1E3F !important; border-bottom-color: #14356E !important; }

/* Auth split-panel layout (Sign Up / Login) */
.auth-left-panel {
    background: linear-gradient(160deg, #14356E 0%, #0B1E3F 100%);
    border-radius: 18px;
    padding: 40px 28px;
    height: 100%;
    color: #FFFFFF;
    text-align: center;

}
.auth-left-panel .auth-icon { font-size: 54px; margin-bottom: 10px; }
.auth-left-panel h2,
.auth-left-panel h2 * { color: #FFFFFF !important; margin-bottom: 4px; font-size: 22px !important; }
.auth-left-panel .auth-tagline { font-size: 26px; font-weight: 800; color: #FFFFFF; margin: 18px 0 10px 0; }
.auth-left-panel .auth-sub { color: #C9C6EE; font-size: 14px; line-height: 1.5; }
.auth-left-panel .auth-emojis { font-size: 30px; margin-top: 24px; letter-spacing: 14px; }
</style>
""", unsafe_allow_html=True)


# =====================================================================================
#  AUTHENTICATION (Sign Up / Email Verification / Login / Forgot Password / Logout)
#  NOTE: all data (users, verification codes, login history) is stored in a local
#  SQLite file. On Streamlit Community Cloud this file resets whenever the app
#  restarts or is redeployed (ephemeral storage). For permanent data across
#  restarts, swap this for an external database (e.g. Supabase/Postgres).
# =====================================================================================
import sqlite3, hashlib, secrets, re, smtplib, ssl
from email.mime.text import MIMEText
from datetime import datetime, timedelta

try:
    import psycopg2
    PSYCOPG2_AVAILABLE = True
except ImportError:
    PSYCOPG2_AVAILABLE = False

DB_PATH = os.path.join(WORK_DIR, "users.db")

# ---- Permanent database (Supabase/Postgres) ----
# If SUPABASE_DB_URL is set in Secrets, all data (users, verification codes,
# login history, worksheet history) is stored permanently in Supabase and
# survives app restarts/redeploys. If not set, falls back to local SQLite
# (temporary — resets on restart), so the app still works without it.
SUPABASE_DB_URL = os.environ.get("SUPABASE_DB_URL", "").strip()
USE_POSTGRES = bool(SUPABASE_DB_URL) and PSYCOPG2_AVAILABLE


class _PGConn:
    """Thin wrapper so Postgres connections support the same .execute()
    convenience method SQLite connections have, and '?' placeholders are
    auto-converted to Postgres's '%s' style."""
    def __init__(self, raw_conn):
        self._conn = raw_conn

    def execute(self, query, params=()):
        cur = self._conn.cursor()
        cur.execute(query.replace("?", "%s"), params)
        return cur

    def commit(self):
        self._conn.commit()

    def rollback(self):
        self._conn.rollback()

    def close(self):
        self._conn.close()


def db_connect():
    if USE_POSTGRES:
        return _PGConn(psycopg2.connect(SUPABASE_DB_URL))
    return sqlite3.connect(DB_PATH)

# ---- Admin / Owner configuration ----
# Set this in Streamlit secrets to decide which username automatically becomes
# the Owner/Admin account (gets access to the Admin Panel with all users +
# full login history). Example in secrets.toml:  ADMIN_USERNAME = "yourname"
ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME", "").strip()

# ---- Email (SMTP) configuration ----
# Needed to actually email verification codes / password-reset codes.
# Example secrets.toml:
#   SMTP_HOST = "smtp.gmail.com"
#   SMTP_PORT = "587"
#   SMTP_USER = "youraddress@gmail.com"
#   SMTP_PASSWORD = "your-gmail-app-password"     (NOT your normal password — use an App Password)
#   SMTP_FROM = "youraddress@gmail.com"
SMTP_HOST = os.environ.get("SMTP_HOST", "").strip()
SMTP_PORT = os.environ.get("SMTP_PORT", "587").strip()
SMTP_USER = os.environ.get("SMTP_USER", "").strip()
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "").strip()
SMTP_FROM = os.environ.get("SMTP_FROM", SMTP_USER).strip()
EMAIL_CONFIGURED = bool(SMTP_HOST and SMTP_USER and SMTP_PASSWORD)


def init_db():
    conn = db_connect()
    autoinc_pk = "SERIAL PRIMARY KEY" if USE_POSTGRES else "INTEGER PRIMARY KEY AUTOINCREMENT"
    conn.execute("""CREATE TABLE IF NOT EXISTS users (
        username TEXT PRIMARY KEY,
        email TEXT UNIQUE,
        password_hash TEXT NOT NULL,
        salt TEXT NOT NULL,
        verified INTEGER DEFAULT 0,
        is_admin INTEGER DEFAULT 0,
        approved INTEGER DEFAULT 0,
        created_at TEXT
    )""")
    conn.commit()
    # Migration for tables created before "approved" existed — ignore if it's already there.
    try:
        conn.execute("ALTER TABLE users ADD COLUMN approved INTEGER DEFAULT 0")
        conn.commit()
    except Exception:
        conn.rollback()
    # Migration for tables created before "created_at" existed — ignore if it's already there.
    try:
        conn.execute("ALTER TABLE users ADD COLUMN created_at TEXT")
        conn.commit()
    except Exception:
        conn.rollback()
    conn.execute(f"""CREATE TABLE IF NOT EXISTS login_logs (
        id {autoinc_pk},
        username TEXT,
        timestamp TEXT,
        success INTEGER
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS verification_codes (
        email TEXT,
        code TEXT,
        purpose TEXT,
        expires_at TEXT
    )""")
    conn.execute(f"""CREATE TABLE IF NOT EXISTS worksheet_logs (
        id {autoinc_pk},
        username TEXT,
        timestamp TEXT,
        grade TEXT,
        subject TEXT,
        topic TEXT,
        export_format TEXT,
        num_questions INTEGER,
        num_copies INTEGER
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS remember_tokens (
        token TEXT PRIMARY KEY,
        username TEXT,
        expires_at TEXT
    )""")
    conn.commit()
    conn.close()

    # Safety net: the Owner account (ADMIN_USERNAME) must never be locked out —
    # make sure it's always marked admin + approved, no matter what happened before.
    if ADMIN_USERNAME:
        conn2 = db_connect()
        try:
            conn2.execute(
                "UPDATE users SET is_admin = 1, approved = 1 WHERE lower(username) = lower(?)",
                (ADMIN_USERNAME,),
            )
            conn2.commit()
        except Exception:
            pass
        finally:
            conn2.close()


def _hash_password(password: str, salt: str) -> str:
    return hashlib.sha256((salt + password).encode()).hexdigest()


def check_password_strength(password: str):
    """Returns (is_ok: bool, message: str)."""
    if len(password) < 8:
        return False, "Password must be at least 8 characters long."
    if not re.search(r"[A-Z]", password):
        return False, "Password must contain at least one UPPERCASE letter."
    if not re.search(r"[a-z]", password):
        return False, "Password must contain at least one lowercase letter."
    if not re.search(r"[0-9]", password):
        return False, "Password must contain at least one number."
    if not re.search(r"""[!@#$%^&*(),.?":{}|<>_\-+=\[\]/\\;'~`]""", password):
        return False, "Password must contain at least one special character (e.g. ! @ # $ %)."
    return True, "Strong password."


def send_email(to_email: str, subject: str, body: str) -> bool:
    """Sends an email via SMTP. If SMTP isn't configured, returns False so the
    caller can fall back to showing the code on-screen (useful for local testing)."""
    if not EMAIL_CONFIGURED:
        return False
    try:
        msg = MIMEText(body)
        msg["Subject"] = subject
        msg["From"] = SMTP_FROM
        msg["To"] = to_email
        context = ssl.create_default_context()
        with smtplib.SMTP(SMTP_HOST, int(SMTP_PORT)) as server:
            server.starttls(context=context)
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.sendmail(SMTP_FROM, to_email, msg.as_string())
        return True
    except Exception as e:
        print("Email send failed:", e)
        return False


def generate_code() -> str:
    return f"{secrets.randbelow(1000000):06d}"


def store_verification_code(email: str, purpose: str) -> str:
    code = generate_code()
    expires = (datetime.utcnow() + timedelta(minutes=15)).isoformat()
    conn = db_connect()
    conn.execute("DELETE FROM verification_codes WHERE email = ? AND purpose = ?", (email, purpose))
    conn.execute("INSERT INTO verification_codes (email, code, purpose, expires_at) VALUES (?, ?, ?, ?)",
                 (email, code, purpose, expires))
    conn.commit()
    conn.close()
    return code


def check_verification_code(email: str, purpose: str, code: str) -> bool:
    conn = db_connect()
    row = conn.execute(
        "SELECT code, expires_at FROM verification_codes WHERE email = ? AND purpose = ?",
        (email, purpose),
    ).fetchone()
    if not row:
        conn.close()
        return False
    stored_code, expires_at = row
    valid = (stored_code == code.strip()) and (datetime.utcnow() <= datetime.fromisoformat(expires_at))
    if valid:
        conn.execute("DELETE FROM verification_codes WHERE email = ? AND purpose = ?", (email, purpose))
        conn.commit()
    conn.close()
    return valid


def create_user(username: str, email: str, password: str):
    salt = secrets.token_hex(16)
    ph = _hash_password(password, salt)
    is_admin = 1 if (ADMIN_USERNAME and username.strip().lower() == ADMIN_USERNAME.lower()) else 0
    approved = 1 if is_admin else 0  # the Owner account never needs to approve itself
    conn = db_connect()
    try:
        conn.execute(
            "INSERT INTO users (username, email, password_hash, salt, verified, is_admin, approved, created_at) VALUES (?, ?, ?, ?, 0, ?, ?, ?)",
            (username.strip(), email.strip().lower(), ph, salt, is_admin, approved, datetime.utcnow().isoformat()),
        )
        conn.commit()
        return True, "Account created successfully."
    except sqlite3.IntegrityError:
        return False, "That username or email is already registered."
    except Exception as e:
        if "unique" in str(e).lower() or "duplicate" in str(e).lower():
            return False, "That username or email is already registered."
        raise
    finally:
        conn.close()


def mark_verified(email: str):
    conn = db_connect()
    conn.execute("UPDATE users SET verified = 1 WHERE email = ?", (email,))
    conn.commit()
    conn.close()


def get_user_by_username(username: str):
    conn = db_connect()
    row = conn.execute(
        "SELECT username, email, password_hash, salt, verified, is_admin, approved FROM users WHERE username = ?",
        (username.strip(),),
    ).fetchone()
    conn.close()
    return row


def update_password(email: str, new_password: str):
    salt = secrets.token_hex(16)
    ph = _hash_password(new_password, salt)
    conn = db_connect()
    conn.execute("UPDATE users SET password_hash = ?, salt = ? WHERE email = ?", (ph, salt, email))
    conn.commit()
    conn.close()


def log_login_attempt(username: str, success: bool):
    conn = db_connect()
    conn.execute(
        "INSERT INTO login_logs (username, timestamp, success) VALUES (?, ?, ?)",
        (username, datetime.utcnow().isoformat(), 1 if success else 0),
    )
    conn.commit()
    conn.close()


def create_remember_token(username: str) -> str:
    token = secrets.token_urlsafe(32)
    expires = (datetime.utcnow() + timedelta(days=30)).isoformat()
    conn = db_connect()
    conn.execute(
        "INSERT INTO remember_tokens (token, username, expires_at) VALUES (?, ?, ?)",
        (token, username, expires),
    )
    conn.commit()
    conn.close()
    return token


def validate_remember_token(token: str):
    """Returns (username, is_admin) if the token is valid and not expired, else None."""
    if not token:
        return None
    conn = db_connect()
    row = conn.execute(
        "SELECT username, expires_at FROM remember_tokens WHERE token = ?", (token,)
    ).fetchone()
    if not row:
        conn.close()
        return None
    username, expires_at = row
    if datetime.utcnow() > datetime.fromisoformat(str(expires_at)):
        conn.execute("DELETE FROM remember_tokens WHERE token = ?", (token,))
        conn.commit()
        conn.close()
        return None
    user_row = conn.execute(
        "SELECT is_admin, verified, approved FROM users WHERE username = ?", (username,)
    ).fetchone()
    conn.close()
    if not user_row or not user_row[1] or not user_row[2]:
        return None  # account was deleted, unverified, or unapproved since the token was issued
    return username, bool(user_row[0])


def delete_remember_token(token: str):
    if not token:
        return
    conn = db_connect()
    conn.execute("DELETE FROM remember_tokens WHERE token = ?", (token,))
    conn.commit()
    conn.close()


def verify_login(username: str, password: str):
    """Returns (ok: bool, message: str, is_admin: bool)."""
    row = get_user_by_username(username)
    if not row:
        log_login_attempt(username, False)
        return False, "Invalid username or password.", False
    _, email, stored_hash, salt, verified, is_admin, approved = row
    if _hash_password(password, salt) != stored_hash:
        log_login_attempt(username, False)
        return False, "Invalid username or password.", False
    if not verified:
        log_login_attempt(username, False)
        return False, "Please verify your email before logging in (check the Verify Email page).", False
    if not approved:
        log_login_attempt(username, False)
        return False, "Your account is verified but is waiting for approval from the Admin. You'll get an email once you're approved.", False
    log_login_attempt(username, True)
    return True, "Logged in successfully.", bool(is_admin)


def get_all_users():
    conn = db_connect()
    rows = conn.execute("SELECT username, email, verified, is_admin, approved FROM users ORDER BY username").fetchall()
    conn.close()
    return rows


def get_all_users_full():
    """Returns every user's COMPLETE information in one place:
    username, email, verified, admin, approved, signup date,
    last successful login, and total worksheets created.
    This is what powers the 'Full User Information' admin view."""

    def _fmt_date(raw):
        """Converts a raw ISO timestamp into a clean, human-readable string.
        Falls back gracefully if the value is missing or malformed."""
        if not raw:
            return "—"
        try:
            dt = datetime.fromisoformat(str(raw))
            return dt.strftime("%b %d, %Y — %I:%M %p")
        except Exception:
            return str(raw)

    conn = db_connect()
    users = conn.execute(
        "SELECT username, email, verified, is_admin, approved, created_at FROM users ORDER BY username"
    ).fetchall()

    # Last successful login per username
    login_rows = conn.execute(
        "SELECT username, MAX(timestamp) FROM login_logs WHERE success = 1 GROUP BY username"
    ).fetchall()
    last_login_map = {u: t for u, t in login_rows}

    # Worksheets created per username
    count_rows = conn.execute(
        "SELECT username, COUNT(*) FROM worksheet_logs GROUP BY username"
    ).fetchall()
    worksheet_count_map = {u: c for u, c in count_rows}

    conn.close()

    full = []
    for username, email, verified, is_admin, approved, created_at in users:
        last_login_raw = last_login_map.get(username)
        full.append({
            "Username": username,
            "Email": email,
            "Verified": "\u2705 Yes" if verified else "\u274c No",
            "Approved": "\u2705 Yes" if approved else "\U0001F7E1 Pending",
            "Admin": "\U0001F451 Yes" if is_admin else "—",
            "Signup Date": _fmt_date(created_at) if created_at else "Before tracking started",
            "Last Login": _fmt_date(last_login_raw) if last_login_raw else "Never logged in",
            "Worksheets Created": worksheet_count_map.get(username, 0),
        })
    return full


def approve_user(username: str, approved: bool):
    conn = db_connect()
    conn.execute("UPDATE users SET approved = ? WHERE username = ?", (1 if approved else 0, username))
    conn.commit()
    row = conn.execute("SELECT email FROM users WHERE username = ?", (username,)).fetchone()
    conn.close()
    if approved and row:
        send_email(
            row[0],
            "Your SchoolSheet Studio account has been approved!",
            f"Hi {username},\n\nGreat news — your account has been approved by the Admin.\n"
            f"You can now log in and start creating worksheets at {APP_NAME}.\n\nWelcome aboard!"
        )


def get_login_logs(limit: int = 200):
    conn = db_connect()
    rows = conn.execute(
        "SELECT username, timestamp, success FROM login_logs ORDER BY id DESC LIMIT ?", (limit,)
    ).fetchall()
    conn.close()
    return rows


# ---- Admin management actions ----
def delete_user(username: str):
    conn = db_connect()
    row = conn.execute("SELECT email FROM users WHERE username = ?", (username,)).fetchone()
    conn.execute("DELETE FROM users WHERE username = ?", (username,))
    if row:
        conn.execute("DELETE FROM verification_codes WHERE email = ?", (row[0],))
    conn.commit()
    conn.close()


def set_user_admin(username: str, is_admin: bool):
    conn = db_connect()
    conn.execute("UPDATE users SET is_admin = ? WHERE username = ?", (1 if is_admin else 0, username))
    conn.commit()
    conn.close()


def set_user_verified(username: str, verified: bool):
    conn = db_connect()
    conn.execute("UPDATE users SET verified = ? WHERE username = ?", (1 if verified else 0, username))
    conn.commit()
    conn.close()


# ---- Worksheet generation logging (for Admin Panel "usage" stats) ----
def log_worksheet_creation(username, grade, subject, topic, export_format, num_questions, num_copies):
    conn = db_connect()
    conn.execute(
        "INSERT INTO worksheet_logs (username, timestamp, grade, subject, topic, export_format, num_questions, num_copies) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
        (username, datetime.utcnow().isoformat(), str(grade), str(subject), str(topic),
         export_format, num_questions, num_copies),
    )
    conn.commit()
    conn.close()


def get_worksheet_logs(limit: int = 200):
    conn = db_connect()
    rows = conn.execute(
        "SELECT username, timestamp, grade, subject, topic, export_format, num_questions, num_copies "
        "FROM worksheet_logs ORDER BY id DESC LIMIT ?", (limit,)
    ).fetchall()
    conn.close()
    return rows


def get_admin_stats():
    conn = db_connect()
    total_users = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    verified_users = conn.execute("SELECT COUNT(*) FROM users WHERE verified = 1").fetchone()[0]
    total_worksheets = conn.execute("SELECT COUNT(*) FROM worksheet_logs").fetchone()[0]
    today = datetime.utcnow().strftime("%Y-%m-%d")
    worksheets_today = conn.execute(
        "SELECT COUNT(*) FROM worksheet_logs WHERE timestamp LIKE ?", (today + "%",)
    ).fetchone()[0]
    failed_logins_today = conn.execute(
        "SELECT COUNT(*) FROM login_logs WHERE success = 0 AND timestamp LIKE ?", (today + "%",)
    ).fetchone()[0]
    conn.close()
    return {
        "total_users": total_users,
        "verified_users": verified_users,
        "total_worksheets": total_worksheets,
        "worksheets_today": worksheets_today,
        "failed_logins_today": failed_logins_today,
    }


init_db()

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = None
if "is_admin" not in st.session_state:
    st.session_state.is_admin = False
if "questions" not in st.session_state:
    st.session_state.questions = []
if "pending_verify_email" not in st.session_state:
    st.session_state.pending_verify_email = None
if "pending_reset_email" not in st.session_state:
    st.session_state.pending_reset_email = None

# ---- "Remember Me" auto-login: if a valid token is in the URL, log the user in
# without asking for username/password again (persists across browser restarts).
if not st.session_state.logged_in:
    _rt = st.query_params.get("rt")
    if _rt:
        _result = validate_remember_token(_rt)
        if _result:
            st.session_state.logged_in = True
            st.session_state.username = _result[0]
            st.session_state.is_admin = _result[1]
        else:
            # stale/expired/invalid token — clean it out of the URL
            st.query_params.pop("rt", None)


# =====================================================================================
#  PAGES: Home / About / Sign Up / Verify Email / Login / Forgot Password / Admin
# =====================================================================================
def home_page():
    _, mid, _ = st.columns([1, 2, 1])
    with mid:
        with st.container(border=True):
            st.markdown(f"<h2 style='text-align:center;'>\U0001F4D8 Welcome to {APP_NAME}</h2>", unsafe_allow_html=True)
            st.markdown(f"""
A professional worksheet generator built for teachers, covering **Grades 1 to 10**.

**What you can do here:**
- Generate curriculum-based or free-content worksheets using AI (Anthropic Claude or Google Gemini)
- Or use your own question bank — completely free, no AI required
- Choose from Multiple Choice, Fill in the Blanks, Match the Column, Short, and Detailed question types
- Export polished, print-ready worksheets as PDF or Word, with your school's branding, logo, and layout

Use the menu on the left to **Sign Up** or **Log In**, then open the **Dashboard** to get started.
""")


def about_page():
    st.title("\u2139\ufe0f About SchoolSheet Studio")
    st.markdown("""
**SchoolSheet Studio** helps teachers create professional, ready-to-print worksheets in minutes.

**Features:**
- Grades 1-10 support with age-appropriate visual themes
- Multiple question types: MCQ, Fill in the Blanks, Match the Column, Short and Detailed Questions
- Three generation methods: curriculum-based, free-content, or your own question bank
- Proper Urdu and English language support
- PDF and Word export with custom branding, logos, and full layout control (spacing, binding gap, copies, randomization)
- Secure accounts with email verification and password reset

Built with a look and feel comparable to school platforms such as Cambridge, APS, and Roots-style institutions.
""")


def _password_field_with_toggle(label: str, key: str):
    """Renders a password text input with a 'Show password' checkbox next to it."""
    show = st.checkbox(f"Show {label.lower()}", key=f"{key}_show")
    return st.text_input(label, type="default" if show else "password", key=key)


def signup_page():
    left, right = st.columns([1, 1.3])
    with left:
        st.markdown(f"""
        <div class="auth-left-panel">
            <div class="auth-icon">\U0001F393</div>
            <h2 style="color:#FFFFFF !important;">{APP_NAME}</h2>
            <div style="color:#C9C6EE; font-size:14px;">Worksheet Generator</div>
            <div class="auth-tagline">Create. Educate.<br>Inspire.</div>
            <div class="auth-sub">Generate professional worksheets for Grades 1-10 in just a few clicks.</div>
            <div class="auth-emojis">\U0001F4DA \U0001F4CB \U0001F331</div>
        </div>
        """, unsafe_allow_html=True)
    with right:
        with st.container(border=True):
            st.markdown("### Create Your Account")
            st.caption(f"Get started with {APP_NAME}")
            if not EMAIL_CONFIGURED:
                st.warning(
                    "Email sending isn't configured yet (no SMTP secrets set), so verification codes "
                    "will be shown on-screen instead of emailed, for testing purposes."
                )
            with st.form("signup_form"):
                username = st.text_input("Username", placeholder="Enter your username")
                email = st.text_input("Email", placeholder="Enter your email address")
                password = _password_field_with_toggle("Password", "signup_password")
                confirm_password = _password_field_with_toggle("Confirm Password", "signup_confirm_password")
                st.caption("Password must be 8+ characters and include an uppercase letter, a lowercase letter, a number, and a special character.")
                submitted = st.form_submit_button("Create Account", type="primary", use_container_width=True)
            st.markdown("<p style='text-align:center;'>Already have an account? Use the <b>Login</b> page from the menu.</p>", unsafe_allow_html=True)
    if submitted:
        if not username or not email or not password or not confirm_password:
            st.warning("Please fill in all fields.")
        elif password != confirm_password:
            st.error("Passwords do not match.")
        else:
            ok_strength, strength_msg = check_password_strength(password)
            if not ok_strength:
                st.error(strength_msg)
            else:
                ok, msg = create_user(username, email, password)
                if ok:
                    code = store_verification_code(email.strip().lower(), "signup")
                    emailed = send_email(
                        email, "Verify your SchoolSheet Studio account",
                        f"Your verification code is: {code}\nIt expires in 15 minutes."
                    )
                    st.session_state.pending_verify_email = email.strip().lower()
                    st.success(msg + " Go to the 'Verify Email' page to activate your account.")
                    if not emailed:
                        st.info(f"(Email not configured — your verification code is: **{code}**)")
                else:
                    st.error(msg)


def verify_email_page():
    st.title("\u2705 Verify Email")
    default_email = st.session_state.pending_verify_email or ""
    email = st.text_input("Email", value=default_email)
    code = st.text_input("Verification Code")

    col1, col2 = st.columns([1, 1])
    with col1:
        verify_clicked = st.button("Verify", type="primary")
    with col2:
        resend_clicked = st.button("\U0001F504 Resend Code")

    if verify_clicked:
        if not email or not code:
            st.warning("Please enter both your email and the verification code.")
        elif check_verification_code(email.strip().lower(), "signup", code):
            mark_verified(email.strip().lower())
            st.success(
                "Email verified successfully! Your account is now waiting for approval from the Admin — "
                "you'll get an email as soon as you're approved, then you can log in."
            )
            st.session_state.pending_verify_email = None
        else:
            st.error("Invalid or expired code. Please try again or click 'Resend Code' below.")

    if resend_clicked:
        if not email:
            st.warning("Please enter your email first, then click Resend Code.")
        else:
            new_code = store_verification_code(email.strip().lower(), "signup")
            emailed = send_email(
                email, "Your new SchoolSheet Studio verification code",
                f"Your new verification code is: {new_code}\nIt expires in 15 minutes."
            )
            if emailed:
                st.success(f"A new verification code has been sent to {email}.")
            else:
                st.success("A new verification code has been generated.")
                st.info(f"(Email not configured — your verification code is: **{new_code}**)")


def login_page():
    left, right = st.columns([1, 1.3])
    with left:
        st.markdown(f"""
        <div class="auth-left-panel">
            <div class="auth-icon">\U0001F393</div>
            <h2 style="color:#FFFFFF !important;">{APP_NAME}</h2>
            <div style="color:#C9C6EE; font-size:14px;">Worksheet Generator</div>
            <div class="auth-tagline">Welcome<br>Back.</div>
            <div class="auth-sub">Log in to continue creating professional worksheets for your class.</div>
            <div class="auth-emojis">\U0001F4DA \U0001F4CB \U0001F331</div>
        </div>
        """, unsafe_allow_html=True)
    with right:
        with st.container(border=True):
            st.markdown("### Log In")
            st.caption(f"Welcome back to {APP_NAME}")
            username = st.text_input("Username", placeholder="Enter your username")
            password = _password_field_with_toggle("Password", "login_password")
            remember_me = st.checkbox("Remember me on this browser", value=True)
            login_clicked = st.button("Log In", type="primary", use_container_width=True)
            st.caption("Forgot your password? Use the 'Forgot Password' page from the menu.")
    if login_clicked:
        if not username or not password:
            st.warning("Please enter both username and password.")
        else:
            ok, msg, is_admin = verify_login(username, password)
            if ok:
                st.session_state.logged_in = True
                st.session_state.username = username
                st.session_state.is_admin = is_admin
                if remember_me:
                    st.query_params["rt"] = create_remember_token(username)
                st.success(msg)
                st.rerun()
            else:
                st.error(msg)


def forgot_password_page():
    st.title("\U0001F513 Forgot Password")
    step = st.radio("Step", ["1. Request Code", "2. Reset Password"], horizontal=True)

    if step == "1. Request Code":
        email = st.text_input("Enter your account email")
        if st.button("Send Reset Code"):
            conn = db_connect()
            row = conn.execute("SELECT username FROM users WHERE email = ?", (email.strip().lower(),)).fetchone()
            conn.close()
            if not row:
                st.error("No account found with that email.")
            else:
                code = store_verification_code(email.strip().lower(), "reset")
                emailed = send_email(
                    email, "SchoolSheet Studio password reset code",
                    f"Your password reset code is: {code}\nIt expires in 15 minutes."
                )
                st.session_state.pending_reset_email = email.strip().lower()
                st.success("Reset code generated. Go to Step 2 to enter it along with your new password.")
                if not emailed:
                    st.info(f"(Email not configured — your reset code is: **{code}**)")

    else:
        default_email = st.session_state.pending_reset_email or ""
        email = st.text_input("Email", value=default_email)
        code = st.text_input("Reset Code")
        new_password = _password_field_with_toggle("New Password", "reset_new_password")
        confirm_password = _password_field_with_toggle("Confirm New Password", "reset_confirm_password")
        if st.button("Reset Password"):
            if not email or not code or not new_password or not confirm_password:
                st.warning("Please fill in all fields.")
            elif new_password != confirm_password:
                st.error("Passwords do not match.")
            else:
                ok_strength, strength_msg = check_password_strength(new_password)
                if not ok_strength:
                    st.error(strength_msg)
                elif check_verification_code(email.strip().lower(), "reset", code):
                    update_password(email.strip().lower(), new_password)
                    st.success("Password reset successfully! You can now log in with your new password.")
                    st.session_state.pending_reset_email = None
                else:
                    st.error("Invalid or expired code.")


def admin_panel_page():
    st.title("\U0001F6E1\ufe0f Admin Panel")
    st.caption("Owner-only view: full control over users, usage, and system status.")

    # ---- System status ----
    st.subheader("\u2699\ufe0f System Status")
    s1, s2, s3, s4 = st.columns(4)
    with s1:
        st.metric("Email (SMTP)", "Configured \u2705" if EMAIL_CONFIGURED else "Not configured \u26a0\ufe0f")
    with s2:
        st.metric("AI Provider", (ACTIVE_PROVIDER.title() if AI_ENABLED else "Off \u26a0\ufe0f"))
    with s3:
        st.metric("Owner Username", ADMIN_USERNAME or "Not set")
    with s4:
        st.metric("Data Storage", "Permanent (Supabase) \u2705" if USE_POSTGRES else "Temporary (resets) \u26a0\ufe0f")

    # ---- Usage stats ----
    st.subheader("\U0001F4CA Usage Overview")
    stats = get_admin_stats()
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Total Users", stats["total_users"])
    c2.metric("Verified Users", stats["verified_users"])
    c3.metric("Worksheets Made", stats["total_worksheets"])
    c4.metric("Worksheets Today", stats["worksheets_today"])
    c5.metric("Failed Logins Today", stats["failed_logins_today"])

    st.divider()

    # ---- User management ----
    st.subheader("\U0001F465 Registered Users")

    users = get_all_users()
    pending = [u for u in users if u[2] and not u[4]]  # verified but not approved
    if pending:
        st.warning(
            f"\u23F3 **{len(pending)} user(s) are verified and waiting for your approval:** "
            + ", ".join(u[0] for u in pending)
        )

    with st.expander("\U0001F4CB Full user list (all info at a glance — check/uncheck boxes and click Apply to save)", expanded=True):
        table_rows = [
            {"Username": u, "Email": e, "Verified": bool(v), "Approved": bool(a), "Admin": bool(ad)}
            for u, e, v, ad, a in users
        ]
        edited = st.data_editor(
            table_rows,
            use_container_width=True, hide_index=True, key="admin_user_editor",
            disabled=["Username", "Email"],
            column_config={
                "Verified": st.column_config.CheckboxColumn(),
                "Approved": st.column_config.CheckboxColumn(),
                "Admin": st.column_config.CheckboxColumn(),
            },
        )
        if st.button("\U0001F4BE Apply Changes", type="primary"):
            changes = 0
            for original, new in zip(table_rows, edited):
                uname = original["Username"]
                if uname == st.session_state.username:
                    continue  # never let the owner accidentally demote/unapprove themselves here
                if original["Verified"] != new["Verified"]:
                    set_user_verified(uname, new["Verified"]); changes += 1
                if original["Approved"] != new["Approved"]:
                    approve_user(uname, new["Approved"]); changes += 1
                if original["Admin"] != new["Admin"]:
                    set_user_admin(uname, new["Admin"]); changes += 1
            if changes:
                st.success(f"Applied {changes} change(s).")
                st.rerun()
            else:
                st.info("No changes to apply.")

    st.divider()

    # ---- Complete user information (read-only, everything in one place) ----
    st.subheader("\U0001F4C4 Complete User Information")
    st.caption("Every user's full profile: signup date, last login, and total worksheets created.")

    full_users = get_all_users_full()
    if full_users:
        full_search = st.text_input("\U0001F50D Search complete info by username or email", key="admin_full_user_search")
        display_rows = full_users
        if full_search:
            fs = full_search.strip().lower()
            display_rows = [
                r for r in full_users
                if fs in r["Username"].lower() or fs in r["Email"].lower()
            ]

        st.dataframe(display_rows, use_container_width=True, hide_index=True)

        csv_data = pd.DataFrame(display_rows).to_csv(index=False).encode("utf-8")
        st.download_button(
            "\U0001F4E5 Download Complete User Info (CSV)",
            csv_data,
            "all_users_complete_info.csv",
            "text/csv",
        )
    else:
        st.info("No users found.")

    st.divider()

    search = st.text_input("\U0001F50D Search by username or email", key="admin_user_search")
    if search:
        s_lower = search.strip().lower()
        users = [u for u in users if s_lower in u[0].lower() or s_lower in u[1].lower()]

    if not users:
        st.info("No users found.")
    else:
        for username, email, verified, is_admin, approved in users:
            status_bits = (
                f"{'\u2705 Verified' if verified else '\u274c Not verified'} \u00b7 "
                f"{'\U0001F7E2 Approved' if approved else '\U0001F7E1 Pending approval'}"
            )
            with st.expander(f"{'\U0001F451 ' if is_admin else ''}{username}  —  {email}  ({status_bits})"):
                col1, col2, col3, col4, col5 = st.columns(5)
                is_self = (username == st.session_state.username)

                with col1:
                    if verified:
                        if st.button("Mark Unverified", key=f"unverify_{username}"):
                            set_user_verified(username, False)
                            st.rerun()
                    else:
                        if st.button("Verify Manually", key=f"verify_{username}"):
                            set_user_verified(username, True)
                            st.success(f"{username} marked as verified.")
                            st.rerun()

                with col2:
                    if approved:
                        if st.button("Revoke Approval", key=f"unapprove_{username}", disabled=is_self):
                            approve_user(username, False)
                            st.rerun()
                    else:
                        if st.button("\u2705 Approve", key=f"approve_{username}", type="primary"):
                            approve_user(username, True)
                            st.success(f"{username} approved — they've been emailed and can now log in.")
                            st.rerun()

                with col3:
                    if is_admin:
                        if st.button("Remove Admin", key=f"demote_{username}", disabled=is_self):
                            set_user_admin(username, False)
                            st.rerun()
                    else:
                        if st.button("Make Admin", key=f"promote_{username}"):
                            set_user_admin(username, True)
                            st.success(f"{username} is now an admin.")
                            st.rerun()

                with col4:
                    confirm_key = f"confirm_delete_{username}"
                    if st.session_state.get(confirm_key):
                        if st.button("\u26a0\ufe0f Confirm Delete", key=f"confirm_del_btn_{username}", disabled=is_self):
                            delete_user(username)
                            st.session_state[confirm_key] = False
                            st.success(f"{username} deleted.")
                            st.rerun()
                    else:
                        if st.button("\U0001F5D1\ufe0f Delete User", key=f"delete_{username}", disabled=is_self):
                            st.session_state[confirm_key] = True
                            st.rerun()

                with col5:
                    if is_self:
                        st.caption("(This is you)")

    st.divider()

    # ---- Worksheet generation history ----
    st.subheader("\U0001F4D8 Worksheet Generation History (most recent first)")
    wlogs = get_worksheet_logs()
    if wlogs:
        st.dataframe(
            [
                {
                    "Username": u, "Timestamp (UTC)": t, "Grade": g, "Subject": subj,
                    "Topic": topic, "Format": fmt, "Questions": nq, "Copies": nc,
                }
                for u, t, g, subj, topic, fmt, nq, nc in wlogs
            ],
            use_container_width=True,
        )
    else:
        st.info("No worksheets generated yet.")

    st.divider()

    # ---- Login history ----
    st.subheader("\U0001F510 Login History (most recent first)")
    logs = get_login_logs()
    if logs:
        st.dataframe(
            [{"Username": u, "Timestamp (UTC)": t, "Success": bool(s)} for u, t, s in logs],
            use_container_width=True,
        )
    else:
        st.info("No login attempts recorded yet.")


def render_dashboard():
    if "questions" not in st.session_state:
        st.session_state.questions = []

    with st.container(border=True):
        st.markdown(f"### \U0001F44B Welcome back, {st.session_state.username}!")
        st.caption(f"{APP_NAME} — create amazing worksheets today.")
        st.markdown(
            "A professional worksheet generator for teachers, covering **Grades 1 to 10**, "
            "using AI (Anthropic Claude or Google Gemini) or your own free question bank."
        )

    # ---- Quick stats for this teacher ----
    _my_logs = [row for row in get_worksheet_logs(limit=1000) if row[0] == st.session_state.username]
    _my_total = len(_my_logs)
    _my_today = sum(1 for row in _my_logs if row[1].startswith(datetime.utcnow().strftime("%Y-%m-%d")))
    _my_questions = sum(row[6] for row in _my_logs) if _my_logs else 0

    m1, m2, m3 = st.columns(3)
    m1.metric("Worksheets Created", _my_total)
    m2.metric("Created Today", _my_today)
    m3.metric("Total Questions Used", _my_questions)

    st.title(f"\U0001F4D8 {APP_NAME}")
    st.caption("Professional Worksheet Generator — Grades 1 to 10")

    if AI_ENABLED:
        _provider_name = "Anthropic (Claude)" if ACTIVE_PROVIDER == "anthropic" else "Google (Gemini)"
        if st.session_state.is_admin:
            st.success(f"AI question generation is **enabled** using **{_provider_name}**.")
    else:
        if st.session_state.is_admin:
            st.warning(
                "AI question generation is **off** — no AI provider key is set. "
                "Add `ANTHROPIC_API_KEY` or `GEMINI_API_KEY` under Settings \u2192 Secrets to enable Methods 1 & 2 for all users. "
                "(Only you see this message, as the Owner/Admin.)"
            )
        else:
            st.info(
                "AI-assisted generation (Methods 1 & 2) isn't available right now — "
                "please use the **\U0001F4D3 Question Bank** page instead."
            )

    st.markdown("#### \U0001F4C4 Recent Worksheets")
    rc1, rc2 = st.columns([4, 1])
    with rc2:
        st.caption(" ")
    if _my_logs:
        st.dataframe(
            [
                {
                    "Grade": g, "Subject": subj, "Topic": topic,
                    "Format": fmt, "Updated On": t[:16].replace("T", " ") + " UTC",
                }
                for u, t, g, subj, topic, fmt, nq, nc in _my_logs[:5]
            ],
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.info("No worksheets yet — start creating from **Create Worksheet** in the sidebar.")



def render_create_worksheet():
    with st.expander("\U0001F9ED Getting Started — Click here if you're new", expanded=True):
        st.markdown("""
    ### \U0001F44B New here? Start with this

    **1. Set your Branding on \U0001F3A8 Cover Designer and your layout/copies on \u2699\ufe0f School Settings** (sidebar) — do this once, it applies to every worksheet.

    **2. Pick Grade / Subject / Topic / Question Types below, then generate questions — pick ONE:**
    - **Method 1 — Curriculum-Based:** Upload a document (PDF/DOCX/TXT) containing content or questions to build the worksheet from it, OR simply generate directly from the Grade/Subject/Topic you selected above. *(Requires an AI provider)*
    - **Method 2 — Free Content:** Paste or upload any raw text (e.g. copied from ChatGPT) — it will be automatically cleaned and turned into a worksheet. *(Requires an AI provider)*
    - **\U0001F4D3 Question Bank (its own page in the sidebar):** Upload a CSV/Excel file of pre-written questions — completely free, no AI required.

    **3. Go to the Preview & Export tab, choose PDF or Word format, and click "Create Worksheet" — a download button will appear immediately.**
    """)

    def _save_uploaded(uploaded_file):
        """Streamlit gives file-like objects, not paths — save to a temp path for the existing helpers."""
        if uploaded_file is None:
            return None
        suffix = os.path.splitext(uploaded_file.name)[1]
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=WORK_DIR)
        tmp.write(uploaded_file.getvalue())
        tmp.close()
        return tmp.name

    st.subheader("\u2699\ufe0f Worksheet Content Settings")
    st.caption(
        "Branding (school name, logo, colors) now lives on **\U0001F3A8 Cover Designer**, and layout/copies "
        "settings live on **\u2699\ufe0f School Settings** — both pages in the sidebar. Adjust them there; "
        "this page will use whatever you've set."
    )

    c1, c2, c3 = st.columns(3)
    with c1:
        grade = st.selectbox("Grade", GRADES, index=2, key="cw_grade")
    with c2:
        subject = st.selectbox("Subject", SUBJECTS, index=0, key="cw_subject")
    with c3:
        language = st.radio("Language", ["English", "Urdu"], horizontal=True, key="cw_language")

    topic = st.text_input("Topic", placeholder="e.g. Nouns, Photosynthesis, Fractions", key="cw_topic")
    q_types = st.multiselect(
        "Question Types (select any combination)", QUESTION_TYPES,
        default=["Choose the Best Option (MCQ)"], key="cw_q_types",
    )
    num_q = st.number_input("Total Questions", min_value=1, value=10, step=1, key="cw_num_q")

    with st.expander("\U0001F4CB Current Branding & Layout Settings (from Cover Designer / School Settings)"):
        st.caption(
            f"School: **{st.session_state.get('cd_school_name', 'Your School Name')}** · "
            f"Theme: **{st.session_state.get('cd_visual_theme', THEMES[0])}** · "
            f"Copies: **{st.session_state.get('ss_num_copies', 1)}** ({st.session_state.get('ss_copy_mode', 'Identical')})"
        )

    def build_cfg():
        logo_file = st.session_state.get("cd_logo_file")
        logo_path = _save_uploaded(logo_file) if logo_file else None
        return WorksheetConfig(
            school_name=st.session_state.get("cd_school_name") or "Your School Name",
            logo_path=logo_path,
            logo_size_pt=60,
            cover_color_hex=st.session_state.get("cd_cover_color", "#2563EB"),
            grade=str(grade),
            subject=subject,
            topic=topic or "General Topic",
            language=language,
            question_types=q_types or ["Choose the Best Option (MCQ)"],
            num_questions=int(num_q),
            questions_per_page=(int(st.session_state.get("ss_q_per_page", 0)) or None),
            spacing_pt=int(st.session_state.get("ss_spacing", 18)),
            binding_gap_mm=int(st.session_state.get("ss_binding_gap", 0)),
            num_copies=int(st.session_state.get("ss_num_copies", 1)),
            copy_mode=st.session_state.get("ss_copy_mode", "Identical"),
            visual_theme=st.session_state.get("cd_visual_theme", THEMES[0]),
            show_student_name=st.session_state.get("ss_show_name", True),
            show_class=st.session_state.get("ss_show_class", True),
            show_roll_no=st.session_state.get("ss_show_roll", True),
            show_date=st.session_state.get("ss_show_date", True),
        )


    def preview_markdown(questions):
        if not questions:
            return "No questions generated yet."
        out = []
        for i, q in enumerate(questions, 1):
            out.append(f"**{i}. ({q.q_type})** {q.text}")
            if q.options:
                out.append("   " + "   ".join(f"({chr(97+j)}) {o}" for j, o in enumerate(q.options)))
            if q.match_left and q.match_right:
                out.append(f"   Match: {q.match_left}  \u2194  {q.match_right}")
            if q.image_keyword:
                out.append(f"   \U0001F5BC illustration: {q.image_keyword}")
            out.append("")
        return "\n".join(out)


    tab1, tab2, tab3 = st.tabs([
        "\U0001F4C2 Method 1: Curriculum-Based",
        "\U0001F4DD Method 2: Free Content",
        "\U0001F441\ufe0f Preview & Export",
    ])

    with tab1:
        st.markdown("Use **either** Path A or Path B.")
        pathA, pathB = st.tabs(["Path A — Upload File + Chapter/Topic/Page", "Path B — Use Grade/Subject/Topic above"])
        with pathA:
            m1a_file = st.file_uploader("Upload Curriculum File (PDF/DOCX/TXT)", key="m1a_file")
            colA1, colA2, colA3 = st.columns(3)
            m1a_chapter = colA1.text_input("Chapter", key="m1a_chapter")
            m1a_topic = colA2.text_input("Topic", key="m1a_topic")
            m1a_page = colA3.text_input("Page No.", key="m1a_page")
            if st.button("Generate from File \u2728", key="m1a_btn", type="primary"):
                if not m1a_file:
                    st.warning("Please upload a curriculum file first.")
                else:
                    try:
                        path = _save_uploaded(m1a_file)
                        raw = extract_text_from_file(path)
                        focus_note = f"\n\n[Focus strictly on: Chapter '{m1a_chapter}', Topic '{m1a_topic}', Page {m1a_page}]"
                        cfg = build_cfg()
                        cfg.topic = m1a_topic or cfg.topic
                        st.session_state.questions = generate_questions_ai(cfg, source_text=raw + focus_note)
                        st.success(f"Generated {len(st.session_state.questions)} questions.")
                    except RuntimeError as e:
                        st.error(str(e))
        with pathB:
            st.markdown("Uses the Grade / Subject / Topic selected above in Worksheet Settings.")
            if st.button("Generate from Selection \u2728", key="m1b_btn", type="primary"):
                try:
                    cfg = build_cfg()
                    st.session_state.questions = generate_questions_ai(cfg, source_text=None)
                    st.success(f"Generated {len(st.session_state.questions)} questions.")
                except RuntimeError as e:
                    st.error(str(e))

    with tab2:
        m2_text = st.text_area("Paste content here (e.g. from ChatGPT, notes, textbook excerpt)", height=200)
        m2_file = st.file_uploader("Or upload a file instead", key="m2_file")
        if st.button("Clean & Generate \u2728", key="m2_btn", type="primary"):
            raw = m2_text or ""
            if m2_file:
                raw += "\n" + extract_text_from_file(_save_uploaded(m2_file))
            if not raw.strip():
                st.warning("Please paste text or upload a file first.")
            else:
                try:
                    cleaned = clean_pasted_content(raw)
                    cfg = build_cfg()
                    st.session_state.questions = generate_questions_ai(cfg, source_text=cleaned)
                    st.success(f"Generated {len(st.session_state.questions)} questions.")
                except RuntimeError as e:
                    st.error(str(e))

    with tab3:
        st.info("Looking for **Method 3 — Your Own Question Bank**? It now has its own page: \U0001F4D3 **Question Bank** in the sidebar.")
        st.markdown(preview_markdown(st.session_state.questions))
        export_format = st.radio("Export Format", ["PDF", "Word"], horizontal=True)
        include_answer_key = st.checkbox("Include a separate Answer Key page at the end", value=True)
        if st.button("\U0001F4C4 Create Worksheet", type="primary"):
            if not st.session_state.questions:
                st.warning("Please generate/preview a worksheet first.")
            else:
                try:
                    cfg = build_cfg()
                    cfg.include_answer_key = include_answer_key
                    out_dir = tempfile.mkdtemp(dir=WORK_DIR)
                    if export_format == "PDF":
                        paths = build_pdf_copies(cfg, st.session_state.questions, out_dir)
                    else:
                        paths = build_docx_copies(cfg, st.session_state.questions, out_dir)

                    log_worksheet_creation(
                        st.session_state.username, cfg.grade, cfg.subject, cfg.topic,
                        export_format, len(st.session_state.questions), len(paths),
                    )

                    if len(paths) == 1:
                        with open(paths[0], "rb") as f:
                            st.download_button("\u2b07\ufe0f Download Worksheet", f, file_name=os.path.basename(paths[0]))
                        st.success(f"Worksheet generated ({export_format}).")
                    else:
                        zip_path = shutil.make_archive(os.path.join(out_dir, "worksheets"), "zip", out_dir)
                        with open(zip_path, "rb") as f:
                            st.download_button("\u2b07\ufe0f Download All Copies (.zip)", f, file_name="worksheets.zip")
                        st.success(f"{len(paths)} {export_format} copies generated and zipped.")
                except Exception as e:
                    st.error(f"Could not generate the worksheet file: {e}")
                    st.exception(e)



def render_my_worksheets():
    st.title("\U0001F4C1 My Worksheets")
    st.caption("Every worksheet you've generated, most recent first.")
    logs = [row for row in get_worksheet_logs(limit=1000) if row[0] == st.session_state.username]
    if not logs:
        st.info("You haven't created any worksheets yet. Go to **Create Worksheet** to make your first one.")
        return
    st.dataframe(
        [
            {
                "Grade": g, "Subject": subj, "Topic": topic,
                "Format": fmt, "Questions": nq, "Copies": nc,
                "Updated On": t[:16].replace("T", " ") + " UTC",
            }
            for u, t, g, subj, topic, fmt, nq, nc in logs
        ],
        use_container_width=True,
    )


def render_cover_designer_page():
    st.title("\U0001F3A8 Cover Designer")
    st.caption("School branding — applies automatically to every worksheet you generate.")
    with st.container(border=True):
        c1, c2 = st.columns(2)
        with c1:
            st.text_input("School Name", value=st.session_state.get("cd_school_name", "Your School Name"), key="cd_school_name")
            st.selectbox(
                "Visual Theme", THEMES,
                index=THEMES.index(st.session_state["cd_visual_theme"]) if st.session_state.get("cd_visual_theme") in THEMES else 0,
                key="cd_visual_theme",
                help="Auto = colorful for young grades, professional for older grades.",
            )
        with c2:
            st.color_picker("Cover Color", value=st.session_state.get("cd_cover_color", "#2563EB"), key="cd_cover_color")
            st.file_uploader("School Logo (optional)", type=["png", "jpg", "jpeg"], key="cd_logo_file")
    st.success("Saved automatically — these settings will be used the next time you open **Create Worksheet**.")


def render_school_settings_page():
    st.title("\u2699\ufe0f School Settings")
    st.caption("Layout, fields shown, and number of copies — applies to every worksheet you generate.")
    with st.container(border=True):
        st.markdown("##### Fields shown on the worksheet")
        c1, c2, c3, c4 = st.columns(4)
        c1.checkbox("Student Name", value=st.session_state.get("ss_show_name", True), key="ss_show_name")
        c2.checkbox("Class", value=st.session_state.get("ss_show_class", True), key="ss_show_class")
        c3.checkbox("Roll No.", value=st.session_state.get("ss_show_roll", True), key="ss_show_roll")
        c4.checkbox("Date", value=st.session_state.get("ss_show_date", True), key="ss_show_date")

        st.markdown("##### Layout")
        c5, c6 = st.columns(2)
        c5.slider("Spacing Between Questions (pt)", 0, 60, st.session_state.get("ss_spacing", 18), step=2, key="ss_spacing")
        c6.slider("Binding Gap (mm)", 0, 40, st.session_state.get("ss_binding_gap", 0), step=2, key="ss_binding_gap")
        st.number_input(
            "Questions per Page (0 = unlimited/auto-flow)", min_value=0,
            value=st.session_state.get("ss_q_per_page", 0), step=1, key="ss_q_per_page",
        )

        st.markdown("##### Copies")
        c7, c8 = st.columns(2)
        c7.number_input("Number of Copies", min_value=1, value=st.session_state.get("ss_num_copies", 1), step=1, key="ss_num_copies")
        c8.radio(
            "Copy Mode", ["Identical", "Randomized"], horizontal=True,
            index=["Identical", "Randomized"].index(st.session_state.get("ss_copy_mode", "Identical")),
            key="ss_copy_mode",
        )
    st.success("Saved automatically — these settings will be used the next time you open **Create Worksheet**.")


def render_question_bank_page():
    st.title("\U0001F4D3 Question Bank")
    st.caption("Bring your own pre-written questions — 100% free, no AI key needed.")

    with st.container(border=True):
        st.markdown("##### 1. Filter your question bank")
        c1, c2, c3 = st.columns(3)
        with c1:
            qb_grade = st.selectbox("Grade", GRADES, index=2, key="qb_grade")
        with c2:
            qb_subject = st.selectbox("Subject", SUBJECTS, index=0, key="qb_subject")
        with c3:
            qb_topic = st.text_input("Topic", placeholder="e.g. Nouns", key="qb_topic")
        qb_qtypes = st.multiselect(
            "Question Types", QUESTION_TYPES, default=["Choose the Best Option (MCQ)"], key="qb_qtypes",
        )
        qb_num_q = st.number_input("Total Questions", min_value=1, value=10, step=1, key="qb_num_q")

        st.markdown("##### 2. Upload your file")
        st.caption(
            "Columns required: `grade, subject, topic, q_type, text, options, match_left, match_right, "
            "answer, image_keyword`. Use `|` to separate multiple values (e.g. `Cat|Dog|Cow|Horse`)."
        )
        qb_file = st.file_uploader("Upload Question Bank (CSV or XLSX)", key="qb_file")
        st.download_button(
            "\u2b07\ufe0f Download CSV Template",
            "grade,subject,topic,q_type,text,options,match_left,match_right,answer,image_keyword\n"
            "3,English,Nouns,mcq,\"Which word is a noun?\",\"Run|Jump|Dog|Fast\",,,Dog,dog\n",
            file_name="question_bank_template.csv", mime="text/csv",
        )

        if st.button("\U0001F5C2\ufe0f Load Questions from Bank \u2728", type="primary"):
            if not qb_file:
                st.warning("Please upload a question bank CSV/Excel file first.")
            else:
                try:
                    tmp_cfg = WorksheetConfig(
                        school_name=st.session_state.get("cd_school_name") or "Your School Name",
                        logo_path=None, logo_size_pt=60,
                        cover_color_hex=st.session_state.get("cd_cover_color", "#2563EB"),
                        grade=str(qb_grade), subject=qb_subject, topic=qb_topic or "General Topic",
                        language=st.session_state.get("cw_language", "English"),
                        question_types=qb_qtypes or ["Choose the Best Option (MCQ)"],
                        num_questions=int(qb_num_q),
                        questions_per_page=(int(st.session_state.get("ss_q_per_page", 0)) or None),
                        spacing_pt=int(st.session_state.get("ss_spacing", 18)),
                        binding_gap_mm=int(st.session_state.get("ss_binding_gap", 0)),
                        num_copies=int(st.session_state.get("ss_num_copies", 1)),
                        copy_mode=st.session_state.get("ss_copy_mode", "Identical"),
                        visual_theme=st.session_state.get("cd_visual_theme", THEMES[0]),
                        show_student_name=st.session_state.get("ss_show_name", True),
                        show_class=st.session_state.get("ss_show_class", True),
                        show_roll_no=st.session_state.get("ss_show_roll", True),
                        show_date=st.session_state.get("ss_show_date", True),
                    )
                    suffix = os.path.splitext(qb_file.name)[1]
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=WORK_DIR)
                    tmp.write(qb_file.getvalue())
                    tmp.close()
                    df = load_question_bank(tmp.name)
                    questions = questions_from_bank(df, tmp_cfg)
                    if not questions:
                        st.warning("No matching questions found for this Grade/Subject/Topic/Question-Type combination.")
                    else:
                        st.session_state.questions = questions
                        # Keep Create Worksheet's settings in sync so Preview & Export uses the same values
                        st.session_state["cw_grade"] = str(qb_grade)
                        st.session_state["cw_subject"] = qb_subject
                        st.session_state["cw_topic"] = qb_topic
                        st.session_state["cw_q_types"] = qb_qtypes
                        st.session_state["cw_num_q"] = int(qb_num_q)
                        st.success(f"Loaded {len(questions)} questions \u2705")
                except Exception as e:
                    st.error(f"Could not read question bank: {e}")

    if st.session_state.get("questions"):
        st.markdown("##### 3. Preview")
        with st.expander(f"\U0001F441\ufe0f {len(st.session_state.questions)} questions loaded — click to preview", expanded=False):
            for i, q in enumerate(st.session_state.questions[:10], 1):
                st.markdown(f"**{i}. ({q.q_type})** {q.text}")
        st.info("Go to **\u270D\ufe0f Create Worksheet \u2192 Preview & Export** to download this as PDF or Word.")


def render_analytics_page():
    st.title("\U0001F4C8 Analytics")
    st.caption("Usage trends across all teachers — owner-only view.")

    logs = get_worksheet_logs(limit=5000)
    if not logs:
        st.info("No worksheets have been generated yet — analytics will appear here once teachers start creating worksheets.")
        return

    df = pd.DataFrame(logs, columns=["username", "timestamp", "grade", "subject", "topic", "format", "num_questions", "num_copies"])
    df["date"] = pd.to_datetime(df["timestamp"]).dt.date

    stats = get_admin_stats()
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Worksheets", stats["total_worksheets"])
    c2.metric("Created Today", stats["worksheets_today"])
    c3.metric("Active Teachers", df["username"].nunique())
    c4.metric("Avg. Questions / Worksheet", round(df["num_questions"].mean(), 1))

    st.markdown("#### \U0001F4C5 Worksheets Created — Last 14 Days")
    last_14 = pd.Timestamp.utcnow().tz_localize(None).normalize() - pd.Timedelta(days=13)
    daily = df[pd.to_datetime(df["date"]) >= last_14].groupby("date").size()
    full_range = pd.date_range(last_14, pd.Timestamp.utcnow().tz_localize(None).normalize())
    daily = daily.reindex(full_range.date, fill_value=0)
    st.bar_chart(daily)

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("#### \U0001F4DA By Subject")
        st.bar_chart(df["subject"].value_counts())
    with col2:
        st.markdown("#### \U0001F393 By Grade")
        st.bar_chart(df["grade"].value_counts().sort_index())

    st.markdown("#### \U0001F465 Top Teachers by Worksheets Created")
    top_users = df["username"].value_counts().head(10)
    st.dataframe(
        [{"Teacher": u, "Worksheets Created": c} for u, c in top_users.items()],
        use_container_width=True, hide_index=True,
    )

    st.markdown("#### \U0001F4C4 Export Format Preference")
    st.bar_chart(df["format"].value_counts())


def render_help_page():
    st.title("\u2753 Help & Support")
    with st.container(border=True):
        st.markdown(f"""
**Need help with {APP_NAME}?**

- For account issues (verification codes, login, password reset), use the **Resend Code** button
  on the Verify Email page, or the **Forgot Password** page from the sidebar.
- For anything else, reach out to: **support.schoolsheetstudio@gmail.com**

**Quick FAQ:**
- *Worksheet not generating?* Make sure an AI provider is configured (ask your Admin), or use
  **Method 3 — Question Bank**, which works with no AI key at all.
- *Didn't get your verification email?* Check Spam/Junk, or click **Resend Code**.
""")



# =====================================================================================
#  SIDEBAR NAVIGATION
# =====================================================================================
with st.sidebar:
    st.markdown(f"## \U0001F393 {APP_NAME}")
    st.caption("Professional Worksheet Generator")
    if st.session_state.logged_in:
        st.success(f"Logged in as **{st.session_state.username}**" + (" (Admin)" if st.session_state.is_admin else ""))

        if "_last_nav_group" not in st.session_state:
            st.session_state["_last_nav_group"] = "workspace"

        def _mark_admin_active():
            st.session_state["_last_nav_group"] = "admin"

        def _mark_workspace_active():
            st.session_state["_last_nav_group"] = "workspace"

        if st.session_state.is_admin:
            st.markdown("###### \U0001F6E1\uFE0F ADMIN")
            admin_nav = st.radio(
                "Admin", ["\U0001F6E1\uFE0F Admin Panel", "\U0001F4C8 Analytics"],
                label_visibility="collapsed", key="admin_nav_radio", on_change=_mark_admin_active,
            )
            st.markdown("###### \U0001F4DA WORKSPACE")
        else:
            admin_nav = None

        nav_options = [
            "\U0001F3E0 Home", "\U0001F4CA Dashboard", "\u270D\uFE0F Create Worksheet",
            "\U0001F4C1 My Worksheets", "\U0001F4D3 Question Bank", "\U0001F3A8 Cover Designer",
            "\u2699\uFE0F School Settings", "\u2139\uFE0F About", "\u2753 Help & Support", "\U0001F6AA Logout",
        ]
        workspace_nav = st.radio(
            "Navigate", nav_options,
            label_visibility="collapsed", key="workspace_nav_radio", on_change=_mark_workspace_active,
        )
        nav = admin_nav if (st.session_state["_last_nav_group"] == "admin" and admin_nav) else workspace_nav
    else:
        st.info("Log in or sign up to use the Dashboard.")
        nav = st.radio(
            "Navigate",
            ["\U0001F4DD Sign Up", "\U0001F511 Login", "\U0001F3E0 Home", "\u2139\uFE0F About",
             "\u2705 Verify Email", "\U0001F511 Forgot Password"],
            label_visibility="collapsed",
        )

if nav == "\U0001F3E0 Home":
    home_page()
elif nav == "\u2139\ufe0f About":
    about_page()
elif nav == "\U0001F511 Login":
    login_page()
elif nav == "\U0001F4DD Sign Up":
    signup_page()
elif nav == "\u2705 Verify Email":
    verify_email_page()
elif nav == "\U0001F511 Forgot Password":
    forgot_password_page()
elif nav == "\U0001F6E1\ufe0f Admin Panel":
    if st.session_state.is_admin:
        admin_panel_page()
    else:
        st.warning("Admin access only.")
elif nav == "\U0001F4C8 Analytics":
    if st.session_state.is_admin:
        render_analytics_page()
    else:
        st.warning("Admin access only.")
elif nav == "\U0001F4CA Dashboard":
    if st.session_state.logged_in:
        render_dashboard()
    else:
        st.warning("Please log in to access the Dashboard.")
elif nav == "\u270D\ufe0f Create Worksheet":
    if st.session_state.logged_in:
        render_create_worksheet()
    else:
        st.warning("Please log in to create a worksheet.")
elif nav == "\U0001F4C1 My Worksheets":
    if st.session_state.logged_in:
        render_my_worksheets()
    else:
        st.warning("Please log in to view your worksheets.")
elif nav == "\U0001F4D3 Question Bank":
    if st.session_state.logged_in:
        render_question_bank_page()
    else:
        st.warning("Please log in to use the Question Bank.")
elif nav == "\U0001F3A8 Cover Designer":
    if st.session_state.logged_in:
        render_cover_designer_page()
    else:
        st.warning("Please log in to use the Cover Designer.")
elif nav == "\u2699\ufe0f School Settings":
    if st.session_state.logged_in:
        render_school_settings_page()
    else:
        st.warning("Please log in to view School Settings.")
elif nav == "\u2753 Help & Support":
    render_help_page()
elif nav == "\U0001F6AA Logout":
    _rt = st.query_params.get("rt")
    if _rt:
        delete_remember_token(_rt)
        st.query_params.pop("rt", None)
    st.session_state.logged_in = False
    st.session_state.username = None
    st.session_state.is_admin = False
    st.success("You have been logged out.")
    st.rerun()
