"""
SchoolSheet Studio — Professional Worksheet Generator (Streamlit version)
Deployable for free on Streamlit Community Cloud (share.streamlit.io) —
no credit card, connects directly to a GitHub repo.
"""
import os
import requests

# ------------------------------------------------------------------
# AI Provider & Keys — configured via Streamlit "Secrets"
# (App settings -> Secrets, in TOML format, e.g.:
#    ANTHROPIC_API_KEY = "sk-ant-..."
#    AI_PROVIDER = "anthropic"
#  Streamlit loads secrets into os.environ automatically at startup
#  when they are also mirrored here, so we read from both.)
# ------------------------------------------------------------------
try:
    import streamlit as st
    for _k in ("ANTHROPIC_API_KEY", "GEMINI_API_KEY", "PIXABAY_API_KEY", "AI_PROVIDER"):
        if _k in st.secrets and not os.environ.get(_k):
            os.environ[_k] = str(st.secrets[_k])
except Exception:
    pass

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "").strip()
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
PIXABAY_API_KEY = os.environ.get("PIXABAY_API_KEY", "").strip()
_provider_pref = os.environ.get("AI_PROVIDER", "").strip().lower()

if _provider_pref == "anthropic" and ANTHROPIC_API_KEY:
    AI_ENABLED = True
    ACTIVE_PROVIDER = "anthropic"
elif _provider_pref == "gemini" and GEMINI_API_KEY:
    AI_ENABLED = True
    ACTIVE_PROVIDER = "gemini"
elif _provider_pref in ("none", ""):
    if ANTHROPIC_API_KEY:
        AI_ENABLED = True
        ACTIVE_PROVIDER = "anthropic"
    elif GEMINI_API_KEY:
        AI_ENABLED = True
        ACTIVE_PROVIDER = "gemini"
    else:
        AI_ENABLED = False
        ACTIVE_PROVIDER = None
else:
    AI_ENABLED = False
    ACTIVE_PROVIDER = None

# ------------------------------------------------------------------
# Urdu Unicode font (Noto Nastaliq Urdu) so Urdu text renders correctly
# in PDF and Word output.
# ------------------------------------------------------------------
FONT_DIR = "./fonts"
os.makedirs(FONT_DIR, exist_ok=True)
URDU_FONT_PATH = os.path.join(FONT_DIR, "NotoNastaliqUrdu-Regular.ttf")

if not os.path.exists(URDU_FONT_PATH):
    font_url = "https://raw.githubusercontent.com/google/fonts/main/ofl/notonastaliqurdu/NotoNastaliqUrdu%5Bwght%5D.ttf"
    try:
        r = requests.get(font_url, timeout=30)
        r.raise_for_status()
        with open(URDU_FONT_PATH, "wb") as f:
            f.write(r.content)
    except Exception:
        URDU_FONT_PATH = None

LATIN_FONT_PATH = None
try:
    import matplotlib
    LATIN_FONT_PATH = os.path.join(os.path.dirname(matplotlib.__file__), "mpl-data", "fonts", "ttf", "DejaVuSans.ttf")
    if not os.path.exists(LATIN_FONT_PATH):
        LATIN_FONT_PATH = None
except Exception:
    LATIN_FONT_PATH = None

import pdfplumber
import docx as docx_lib

# ------------------------------------------------------------------
# Unified AI call layer — one function, works with EITHER provider.
# Every other step in this notebook calls _call_ai() and never talks
# to the Anthropic or Gemini SDKs directly, so switching providers in
# Step 2 doesn't require touching any code below.
# ------------------------------------------------------------------
_anthropic_client = None
_gemini_model = None

CLAUDE_MODEL = "claude-sonnet-4-6"
GEMINI_MODEL = "gemini-2.5-flash"

if AI_ENABLED and ACTIVE_PROVIDER == "anthropic":
    from anthropic import Anthropic
    _anthropic_client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
elif AI_ENABLED and ACTIVE_PROVIDER == "gemini":
    import google.generativeai as genai
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    _gemini_model = genai.GenerativeModel(GEMINI_MODEL)


def _call_ai(system_prompt: str, user_content: str, max_tokens: int = 4000) -> str:
    """
    Single entry point for every AI call in this notebook.
    Sends `user_content` with `system_prompt` as instructions and returns
    the raw text reply — regardless of which provider is active.
    Raises RuntimeError if AI is not enabled (no valid key/provider).
    """
    if not AI_ENABLED:
        raise RuntimeError(
            "This feature isn't available right now — please use 'Method 3: Your Own Question Bank' instead."
        )
    if ACTIVE_PROVIDER == "anthropic":
        resp = _anthropic_client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=max_tokens,
            system=system_prompt,
            messages=[{"role": "user", "content": user_content}],
        )
        return "".join(b.text for b in resp.content if b.type == "text").strip()
    elif ACTIVE_PROVIDER == "gemini":
        full_prompt = f"{system_prompt}\n\n{user_content}"
        resp = _gemini_model.generate_content(
            full_prompt,
            generation_config={"max_output_tokens": max_tokens},
        )
        return (getattr(resp, "text", None) or "").strip()
    else:
        raise RuntimeError(f"Unknown AI provider: {ACTIVE_PROVIDER}")


def extract_text_from_file(file_path: str) -> str:
    """Reads .pdf, .docx or .txt and returns raw text."""
    if file_path is None:
        return ""
    lower = file_path.lower()
    try:
        if lower.endswith(".pdf"):
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    text_parts.append(t)
            return "\n".join(text_parts)
        elif lower.endswith(".docx"):
            d = docx_lib.Document(file_path)
            return "\n".join(p.text for p in d.paragraphs)
        else:  # .txt or anything else — read as plain text
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        return f"[ERROR READING FILE: {e}]"


def clean_pasted_content(raw_text: str) -> str:
    """
    Removes AI chatter / preambles / postambles / instructions that are
    NOT actual question content (e.g. "Here are your questions:",
    "I hope this helps!", numbering artifacts from chat exports, etc.)
    Keeps only the real educational content, unaltered.
    """
    if not raw_text or not raw_text.strip():
        return ""

    if not AI_ENABLED:
        # Simple rule-based fallback cleaning when no API key is configured:
        # drops common AI-chatter lines and blank/short filler lines.
        junk_markers = [
            "here are your questions", "here's your worksheet", "i hope this helps",
            "hope this helps", "let me know if", "feel free to", "sure, here",
            "certainly!", "as requested", "best regards", "regards,",
        ]
        lines = raw_text.splitlines()
        kept = []
        for line in lines:
            l = line.strip()
            if not l:
                continue
            if any(marker in l.lower() for marker in junk_markers):
                continue
            kept.append(line)
        return "\n".join(kept).strip()

    system_prompt = (
        "You are a strict text-cleaning engine for a school worksheet system. "
        "You will be given raw text that may contain educational questions or "
        "content mixed with irrelevant chatter (greetings, meta-comments like "
        "'Here are your questions', 'I hope this helps', signatures, extra "
        "headers, disclaimers, or duplicate/broken lines). "
        "Return ONLY the genuine educational content — questions, passages, "
        "topics, instructions meant for students — exactly as originally "
        "written, with no rephrasing. Remove every non-content line. "
        "Do not add any commentary, notes, or explanation of your own. "
        "Output nothing except the cleaned content."
    )

    return _call_ai(system_prompt, raw_text, max_tokens=4000)


print("Text extraction & cleaning utilities ready. Active AI provider:", ACTIVE_PROVIDER or "None (Question Bank mode)")

from dataclasses import dataclass, field
from typing import List, Optional
import json, random

QUESTION_TYPES = [
    "Choose the Best Option (MCQ)",
    "Fill in the Blanks",
    "Match the Column",
    "Short Question",
    "Detailed Question",
]

GRADES = [str(g) for g in range(1, 11)]  # Grade 1 to Grade 10

THEMES = ["Auto (based on grade)", "Colorful / Playful", "Professional / Formal"]


@dataclass
class QuestionItem:
    q_type: str
    text: str
    options: Optional[List[str]] = None       # for MCQ
    match_left: Optional[List[str]] = None     # for Match the Column
    match_right: Optional[List[str]] = None
    answer: Optional[str] = None
    image_keyword: Optional[str] = None        # e.g. "apple", "cat" for young grades

    def to_dict(self):
        return self.__dict__


@dataclass
class WorksheetConfig:
    school_name: str = "Your School Name"
    logo_path: Optional[str] = None
    logo_size_pt: int = 60          # logo width in points
    cover_color_hex: str = "#2563EB"
    font_theme: str = "Poppins/Default"

    grade: str = "3"
    subject: str = "General"
    topic: str = "General Topic"
    language: str = "English"       # "English" or "Urdu"

    question_types: List[str] = field(default_factory=lambda: ["Choose the Best Option (MCQ)"])
    num_questions: int = 10

    questions_per_page: Optional[int] = None   # None = unlimited/auto-flow
    spacing_pt: int = 18                        # vertical space between questions
    binding_gap_mm: int = 0                      # extra margin for binding
    num_copies: int = 1
    copy_mode: str = "Identical"                 # "Identical" or "Randomized"

    visual_theme: str = "Auto (based on grade)"

    show_student_name: bool = True
    show_class: bool = True
    show_roll_no: bool = True
    show_date: bool = True
    include_answer_key: bool = True

    def effective_theme(self) -> str:
        if self.visual_theme != "Auto (based on grade)":
            return self.visual_theme
        return "Colorful / Playful" if int(self.grade) <= 4 else "Professional / Formal"


print("Data models ready. Question types:", QUESTION_TYPES)

import pandas as pd

def _build_generation_prompt(cfg: "WorksheetConfig", source_text: Optional[str]) -> str:
    types_str = ", ".join(cfg.question_types)
    lang_line = (
        "Write the ENTIRE worksheet content in proper, correct, natural Urdu "
        "(Urdu script, not Roman Urdu)."
        if cfg.language == "Urdu"
        else "Write everything in clear, simple, grade-appropriate English."
    )
    source_block = (
        f"Base every question strictly on the following source content "
        f"(do not invent facts outside it):\n\n{source_text}\n"
        if source_text
        else f"Topic: {cfg.topic}\nSubject: {cfg.subject}\n"
    )
    image_line = (
        "Since this is for a young grade (1-4), for at least half the "
        "questions add a simple, common, real-world 'image_keyword' "
        "(e.g. apple, cat, monkey, ball, sun, book) that a teacher could "
        "illustrate the question with. Use only simple everyday objects/animals."
        if int(cfg.grade) <= 4
        else "This is for an older grade — do not add image_keyword, keep it null."
    )

    return f"""
You are an expert school curriculum question-writer for Grade {cfg.grade}, subject {cfg.subject}.
{source_block}
Generate exactly {cfg.num_questions} TOTAL questions, using ONLY these question types: {types_str}.
Distribute the questions across the selected types as evenly as possible.
{lang_line}
{image_line}

Rules:
- Every question must be completely different in content from every other question (no repeats, no rephrasing of the same fact twice).
- For "Choose the Best Option (MCQ)": provide exactly 4 options and mark the correct answer.
- For "Fill in the Blanks": use a blank shown as "_____" inside the sentence, and give the answer.
- For "Match the Column": provide a left list and right list (same length, 4-6 pairs), shuffle the right list order, and give the answer as pairs.
- For "Short Question": a question needing a 1-3 sentence answer.
- For "Detailed Question": a question needing a paragraph-length answer.
- Keep difficulty appropriate for Grade {cfg.grade}.

Return ONLY valid JSON (no markdown fences, no commentary), as a list of objects with this exact schema:
[
  {{
    "q_type": "one of: Choose the Best Option (MCQ) | Fill in the Blanks | Match the Column | Short Question | Detailed Question",
    "text": "the question text (or sentence with blank)",
    "options": ["...", "...", "...", "..."] or null,
    "match_left": ["..."] or null,
    "match_right": ["..."] or null,
    "answer": "the correct answer / answer key",
    "image_keyword": "simple keyword" or null
  }}
]
""".strip()


def generate_questions_ai(cfg: "WorksheetConfig", source_text: Optional[str] = None) -> List[QuestionItem]:
    """Calls the active AI provider (Anthropic or Gemini) to generate a fresh, non-repeating question set."""
    if not AI_ENABLED:
        raise RuntimeError(
            "This feature isn't available right now — please use 'Method 3: Your Own Question Bank' instead."
        )
    prompt = _build_generation_prompt(cfg, source_text)
    raw = _call_ai(
        system_prompt=(
            "You are an expert school curriculum question-writer. "
            "Return ONLY valid JSON (no markdown fences, no commentary)."
        ),
        user_content=prompt,
        max_tokens=4000,
    )
    raw = raw.strip("`")
    if raw.lower().startswith("json"):
        raw = raw[4:].strip()
    try:
        data = json.loads(raw)
    except Exception:
        # last-resort: try to locate the JSON array inside the text
        start = raw.find("[")
        end = raw.rfind("]")
        data = json.loads(raw[start:end + 1])

    return [QuestionItem(
        q_type=item.get("q_type", "Short Question"),
        text=item.get("text", ""),
        options=item.get("options"),
        match_left=item.get("match_left"),
        match_right=item.get("match_right"),
        answer=item.get("answer"),
        image_keyword=item.get("image_keyword"),
    ) for item in data]


def load_question_bank(file_path: str) -> pd.DataFrame:
    """
    Optional plug-in question bank. Expected columns (csv or xlsx):
    grade, subject, topic, q_type, text, options, match_left, match_right, answer, image_keyword
    'options' / 'match_left' / 'match_right' cells should be pipe-separated, e.g.  "A|B|C|D"
    """
    if file_path.lower().endswith(".csv"):
        return pd.read_csv(file_path)
    return pd.read_excel(file_path)


def questions_from_bank(df: "pd.DataFrame", cfg: "WorksheetConfig") -> List[QuestionItem]:
    filtered = df[
        (df["grade"].astype(str) == str(cfg.grade))
        & (df["subject"].str.lower() == cfg.subject.lower())
        & (df["topic"].str.lower() == cfg.topic.lower())
        & (df["q_type"].isin(cfg.question_types))
    ]
    filtered = filtered.sample(frac=1).head(cfg.num_questions)  # shuffle + limit
    items = []
    for _, row in filtered.iterrows():
        split = lambda v: str(v).split("|") if pd.notna(v) else None
        items.append(QuestionItem(
            q_type=row["q_type"],
            text=row["text"],
            options=split(row.get("options")),
            match_left=split(row.get("match_left")),
            match_right=split(row.get("match_right")),
            answer=row.get("answer"),
            image_keyword=row.get("image_keyword") if pd.notna(row.get("image_keyword")) else None,
        ))
    return items


def randomize_copy(questions: List[QuestionItem]) -> List[QuestionItem]:
    """Produces a shuffled variant (question order + MCQ option order) for randomized copies."""
    import copy as _copy
    shuffled = _copy.deepcopy(questions)
    random.shuffle(shuffled)
    for q in shuffled:
        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            correct = q.options[0] if q.answer == q.options[0] else q.answer
            opts = q.options[:]
            random.shuffle(opts)
            q.options = opts
    return shuffled


print("AI generation engine + question bank plug-in ready.")

import hashlib

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_OK = True
except Exception as e:
    print("⚠️ Pillow (PIL) failed to import:", e)
    print("   Images will be skipped, but the rest of the app (PDF/Word generation,")
    print("   AI question generation, all worksheet features) will still work fine.")
    print("   To fix images: Runtime -> Restart session, then run all cells again from Step 1.")
    PIL_OK = False

IMAGE_CACHE_DIR = "./wg_images"
os.makedirs(IMAGE_CACHE_DIR, exist_ok=True)

_FALLBACK_COLORS = [
    "#FF6B6B", "#4ECDC4", "#FFD93D", "#6BCB77", "#4D96FF", "#FF922B", "#C084FC"
]


def _cache_path(keyword: str) -> str:
    h = hashlib.md5(keyword.encode()).hexdigest()[:10]
    return os.path.join(IMAGE_CACHE_DIR, f"{h}.png")


def _make_fallback_icon(keyword: str, path: str, size: int = 300):
    """Generates a clean, simple flat-color icon with the keyword label,
    used only if no Pixabay key / no internet result is available."""
    if not PIL_OK:
        return None
    color = _FALLBACK_COLORS[hash(keyword) % len(_FALLBACK_COLORS)]
    img = Image.new("RGB", (size, size), "white")
    draw = ImageDraw.Draw(img)
    draw.ellipse([20, 20, size - 20, size - 20], fill=color)
    try:
        font = ImageFont.truetype(LATIN_FONT_PATH, 28) if LATIN_FONT_PATH else ImageFont.load_default()
    except Exception:
        font = ImageFont.load_default()
    text = keyword.upper()
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((size - tw) / 2, size - 55), text, fill="white", font=font)
    img.save(path)
    return path


def get_image_for_keyword(keyword: str) -> Optional[str]:
    """Returns a local file path to an image representing `keyword`.
    Tries Pixabay (real photo) first if a key is set, else falls back
    to a clean generated icon so the worksheet is never left blank.
    Returns None (image simply skipped) if Pillow is unavailable."""
    if not keyword:
        return None
    path = _cache_path(keyword)
    if os.path.exists(path):
        return path

    pixabay_key = os.environ.get("PIXABAY_API_KEY", "").strip()
    if pixabay_key:
        try:
            r = requests.get(
                "https://pixabay.com/api/",
                params={
                    "key": pixabay_key,
                    "q": keyword,
                    "image_type": "photo",
                    "category": "animals,food,education",
                    "safesearch": "true",
                    "per_page": 3,
                },
                timeout=15,
            )
            data = r.json()
            if data.get("hits"):
                img_url = data["hits"][0]["webformatURL"]
                img_data = requests.get(img_url, timeout=15).content
                with open(path, "wb") as f:
                    f.write(img_data)
                return path
        except Exception as e:
            print(f"Pixabay fetch failed for '{keyword}', using fallback icon. ({e})")

    return _make_fallback_icon(keyword, path)


print("Image engine ready (real photos via Pixabay, with automatic icon fallback).")

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display

# Register fonts once
_FONTS_REGISTERED = False

def _register_fonts():
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    if URDU_FONT_PATH and os.path.exists(URDU_FONT_PATH):
        pdfmetrics.registerFont(TTFont("Urdu", URDU_FONT_PATH))
    if LATIN_FONT_PATH and os.path.exists(LATIN_FONT_PATH):
        pdfmetrics.registerFont(TTFont("Latin", LATIN_FONT_PATH))
    _FONTS_REGISTERED = True

def _shape(text: str, is_urdu: bool) -> str:
    if not is_urdu:
        return text
    reshaped = arabic_reshaper.reshape(text)
    return get_display(reshaped)


THEME_COLORS = {
    "Colorful / Playful": {"primary": "#FF6B6B", "secondary": "#FFD93D", "accent": "#4ECDC4"},
    "Professional / Formal": {"primary": "#1E3A8A", "secondary": "#334155", "accent": "#2563EB"},
}


def render_pdf(cfg: "WorksheetConfig", questions: List[QuestionItem], output_path: str):
    _register_fonts()
    is_urdu = cfg.language == "Urdu"
    body_font = "Urdu" if (is_urdu and "Urdu" in pdfmetrics.getRegisteredFontNames()) else "Helvetica"
    head_font = "Latin" if "Latin" in pdfmetrics.getRegisteredFontNames() else "Helvetica-Bold"

    theme = THEME_COLORS[cfg.effective_theme()]
    page_w, page_h = A4
    left_margin = 18 * mm + cfg.binding_gap_mm * mm
    right_margin = 18 * mm
    top_margin = 15 * mm
    bottom_margin = 15 * mm

    c = canvas.Canvas(output_path, pagesize=A4)

    def draw_header(page_no_label=""):
        y = page_h - top_margin
        # Logo
        if cfg.logo_path and os.path.exists(cfg.logo_path):
            try:
                c.drawImage(cfg.logo_path, left_margin, y - cfg.logo_size_pt,
                            width=cfg.logo_size_pt, height=cfg.logo_size_pt,
                            preserveAspectRatio=True, mask="auto")
            except Exception:
                pass
        # School name
        c.setFont(head_font, 16)
        c.setFillColor(theme["primary"])
        c.drawCentredString(page_w / 2, y - 20, _shape(cfg.school_name, False))
        c.setFont(head_font, 11)
        c.setFillColor("#333333")
        c.drawCentredString(page_w / 2, y - 38, f"Grade {cfg.grade}  |  {cfg.subject}  |  {cfg.topic}")

        # divider line
        c.setStrokeColor(theme["accent"])
        c.setLineWidth(1.5)
        c.line(left_margin, y - 48, page_w - right_margin, y - 48)

        # student info fields
        fy = y - 65
        c.setFont("Helvetica", 10)
        c.setFillColor("#000000")
        fields = []
        if cfg.show_student_name: fields.append("Student Name: ______________________")
        if cfg.show_class: fields.append("Class: _______")
        if cfg.show_roll_no: fields.append("Roll No: _______")
        if cfg.show_date: fields.append("Date: _______")
        line1 = "     ".join(fields[:2])
        line2 = "     ".join(fields[2:])
        c.drawString(left_margin, fy, line1)
        if line2:
            c.drawString(left_margin, fy - 14, line2)

        c.setStrokeColor("#999999")
        c.setLineWidth(0.7)
        c.line(left_margin, fy - 24, page_w - right_margin, fy - 24)
        return fy - 40  # y position where question body should start

    def new_page():
        c.showPage()
        return draw_header()

    y_cursor = draw_header()
    q_count_on_page = 0
    max_width = page_w - left_margin - right_margin

    def wrap_text(text, font, size, max_w):
        c.setFont(font, size)
        words = text.split(" ")
        lines, cur = [], ""
        for w in words:
            trial = (cur + " " + w).strip()
            if c.stringWidth(trial, font, size) <= max_w:
                cur = trial
            else:
                lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
        return lines

    for idx, q in enumerate(questions, start=1):
        needed_space = 70  # rough estimate; grows below per type
        if y_cursor < bottom_margin + needed_space or (
            cfg.questions_per_page and q_count_on_page >= cfg.questions_per_page
        ):
            y_cursor = new_page()
            q_count_on_page = 0

        c.setFont(body_font, 11)
        c.setFillColor("#000000")
        prefix = f"{idx}. "
        q_text = _shape(prefix + q.text, is_urdu)
        lines = wrap_text(q_text, body_font, 11, max_width)
        for line in lines:
            c.drawString(left_margin, y_cursor, line)
            y_cursor -= 15

        # image for young grades
        if q.image_keyword:
            img_path = get_image_for_keyword(q.image_keyword)
            if img_path:
                try:
                    c.drawImage(img_path, left_margin, y_cursor - 60, width=55, height=55,
                                preserveAspectRatio=True, mask="auto")
                    y_cursor -= 65
                except Exception:
                    pass

        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            for i, opt in enumerate(q.options):
                letter = chr(97 + i)
                c.drawString(left_margin + 15, y_cursor, _shape(f"({letter}) {opt}", is_urdu))
                y_cursor -= 14

        elif q.q_type == "Match the Column" and q.match_left and q.match_right:
            col2_x = left_margin + max_width / 2
            rows = max(len(q.match_left), len(q.match_right))
            for i in range(rows):
                left_txt = q.match_left[i] if i < len(q.match_left) else ""
                right_txt = q.match_right[i] if i < len(q.match_right) else ""
                c.drawString(left_margin + 15, y_cursor, _shape(f"{i+1}. {left_txt}", is_urdu))
                c.drawString(col2_x, y_cursor, _shape(f"{chr(97+i)}. {right_txt}", is_urdu))
                y_cursor -= 14

        elif q.q_type == "Short Question":
            for _ in range(2):
                y_cursor -= 14
                c.setStrokeColor("#cccccc")
                c.line(left_margin, y_cursor, page_w - right_margin, y_cursor)

        elif q.q_type == "Detailed Question":
            for _ in range(5):
                y_cursor -= 14
                c.setStrokeColor("#cccccc")
                c.line(left_margin, y_cursor, page_w - right_margin, y_cursor)

        elif q.q_type == "Fill in the Blanks":
            pass  # blank already embedded in text

        y_cursor -= cfg.spacing_pt
        q_count_on_page += 1

    # ---- Answer Key (separate page at the end) ----
    if getattr(cfg, "include_answer_key", True):
        c.showPage()
        y_cursor = draw_header()
        c.setFont(head_font, 14)
        c.setFillColor(theme["primary"])
        c.drawCentredString(page_w / 2, y_cursor + 10, "Answer Key")
        y_cursor -= 15
        c.setFont(body_font, 11)
        c.setFillColor("#000000")
        for idx, q in enumerate(questions, start=1):
            if y_cursor < bottom_margin + 20:
                y_cursor = new_page()
            ans_text = q.answer if q.answer else "(open-ended — no fixed answer)"
            line = _shape(f"{idx}. {ans_text}", is_urdu)
            for wrapped in wrap_text(line, body_font, 11, max_width):
                c.drawString(left_margin, y_cursor, wrapped)
                y_cursor -= 15

    c.save()
    return output_path


def build_pdf_copies(cfg: "WorksheetConfig", base_questions: List[QuestionItem], output_dir: str) -> List[str]:
    os.makedirs(output_dir, exist_ok=True)
    paths = []
    for i in range(1, cfg.num_copies + 1):
        qs = randomize_copy(base_questions) if cfg.copy_mode == "Randomized" else base_questions
        out = os.path.join(output_dir, f"worksheet_copy_{i}.pdf")
        render_pdf(cfg, qs, out)
        paths.append(out)
    return paths


print("PDF renderer ready.")

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _apply_urdu_font(run, size=12):
    run.font.name = "Noto Nastaliq Urdu"
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), "Noto Nastaliq Urdu")


def render_docx(cfg: "WorksheetConfig", questions: List[QuestionItem], output_path: str):
    is_urdu = cfg.language == "Urdu"
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Mm(18 + cfg.binding_gap_mm)
    section.right_margin = Mm(18)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    theme_primary = RGBColor(0x1E, 0x3A, 0x8A) if cfg.effective_theme() == "Professional / Formal" else RGBColor(0xE0, 0x50, 0x50)

    # Logo + School name
    if cfg.logo_path and os.path.exists(cfg.logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(cfg.logo_path, width=Pt(cfg.logo_size_pt))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(cfg.school_name)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = theme_primary

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Grade {cfg.grade}  |  {cfg.subject}  |  {cfg.topic}").font.size = Pt(11)

    doc.add_paragraph("_" * 90)

    # Student fields
    fields = []
    if cfg.show_student_name: fields.append("Student Name: ______________________")
    if cfg.show_class: fields.append("Class: _______")
    if cfg.show_roll_no: fields.append("Roll No: _______")
    if cfg.show_date: fields.append("Date: _______")
    if fields:
        doc.add_paragraph("     ".join(fields[:2]))
        if len(fields) > 2:
            doc.add_paragraph("     ".join(fields[2:]))
    doc.add_paragraph("_" * 90)

    for idx, q in enumerate(questions, start=1):
        p = doc.add_paragraph()
        if is_urdu:
            _set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{idx}. {q.text}")
        run.bold = True
        if is_urdu:
            _apply_urdu_font(run, 13)
        else:
            run.font.size = Pt(12)

        if q.image_keyword:
            img_path = get_image_for_keyword(q.image_keyword)
            if img_path:
                try:
                    doc.add_picture(img_path, width=Inches(0.7))
                except Exception:
                    pass

        if q.q_type == "Choose the Best Option (MCQ)" and q.options:
            for i, opt in enumerate(q.options):
                op = doc.add_paragraph(f"    ({chr(97+i)}) {opt}")
                if is_urdu:
                    _set_rtl(op)
                    op.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    _apply_urdu_font(op.runs[0], 12)

        elif q.q_type == "Match the Column" and q.match_left and q.match_right:
            table = doc.add_table(rows=max(len(q.match_left), len(q.match_right)), cols=2)
            table.style = "Table Grid"
            for i in range(len(table.rows)):
                l = q.match_left[i] if i < len(q.match_left) else ""
                r_ = q.match_right[i] if i < len(q.match_right) else ""
                table.cell(i, 0).text = f"{i+1}. {l}"
                table.cell(i, 1).text = f"{chr(97+i)}. {r_}"

        elif q.q_type == "Short Question":
            for _ in range(2):
                doc.add_paragraph("_" * 90)

        elif q.q_type == "Detailed Question":
            for _ in range(5):
                doc.add_paragraph("_" * 90)

        # spacing between questions
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(cfg.spacing_pt)

        if cfg.questions_per_page and idx % cfg.questions_per_page == 0 and idx != len(questions):
            doc.add_page_break()

    if getattr(cfg, "include_answer_key", True):
        doc.add_page_break()
        ak_title = doc.add_paragraph()
        ak_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ak_run = ak_title.add_run("Answer Key")
        ak_run.bold = True
        ak_run.font.size = Pt(16)
        ak_run.font.color.rgb = theme_primary

        for idx, q in enumerate(questions, start=1):
            ans_text = q.answer if q.answer else "(open-ended — no fixed answer)"
            p = doc.add_paragraph()
            if is_urdu:
                _set_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"{idx}. {ans_text}")
            if is_urdu:
                _apply_urdu_font(run, 12)
            else:
                run.font.size = Pt(11)

    doc.save(output_path)
    return output_path


def build_docx_copies(cfg: "WorksheetConfig", base_questions: List[QuestionItem], output_dir: str) -> List[str]:
    os.makedirs(output_dir, exist_ok=True)
    paths = []
    for i in range(1, cfg.num_copies + 1):
        qs = randomize_copy(base_questions) if cfg.copy_mode == "Randomized" else base_questions
        out = os.path.join(output_dir, f"worksheet_copy_{i}.docx")
        render_docx(cfg, qs, out)
        paths.append(out)
    return paths


print("DOCX renderer ready.")

# =====================================================================================
#  STREAMLIT UI
# =====================================================================================
import streamlit as st
import tempfile, shutil

APP_NAME = "SchoolSheet Studio"
WORK_DIR = "./wg_output"
os.makedirs(WORK_DIR, exist_ok=True)

SUBJECTS = ["English", "Urdu", "Mathematics", "Science", "Social Studies", "Islamiyat", "Computer Science", "General Knowledge"]

st.set_page_config(page_title=APP_NAME, page_icon="\U0001F4D8", layout="wide")


# =====================================================================================
#  AUTHENTICATION (Sign Up / Email Verification / Login / Forgot Password / Logout)
#  NOTE: all data (users, verification codes, login history) is stored in a local
#  SQLite file. On Streamlit Community Cloud this file resets whenever the app
#  restarts or is redeployed (ephemeral storage). For permanent data across
#  restarts, swap this for an external database (e.g. Supabase/Postgres).
# =====================================================================================
import sqlite3, hashlib, secrets, re, smtplib, ssl
from email.mime.text import MIMEText
from datetime import datetime, timedelta

DB_PATH = os.path.join(WORK_DIR, "users.db")

# ---- Admin / Owner configuration ----
# Set this in Streamlit secrets to decide which username automatically becomes
# the Owner/Admin account (gets access to the Admin Panel with all users +
# full login history). Example in secrets.toml:  ADMIN_USERNAME = "yourname"
ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME", "").strip()

# ---- Email (SMTP) configuration ----
# Needed to actually email verification codes / password-reset codes.
# Example secrets.toml:
#   SMTP_HOST = "smtp.gmail.com"
#   SMTP_PORT = "587"
#   SMTP_USER = "youraddress@gmail.com"
#   SMTP_PASSWORD = "your-gmail-app-password"     (NOT your normal password — use an App Password)
#   SMTP_FROM = "youraddress@gmail.com"
SMTP_HOST = os.environ.get("SMTP_HOST", "").strip()
SMTP_PORT = os.environ.get("SMTP_PORT", "587").strip()
SMTP_USER = os.environ.get("SMTP_USER", "").strip()
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "").strip()
SMTP_FROM = os.environ.get("SMTP_FROM", SMTP_USER).strip()
EMAIL_CONFIGURED = bool(SMTP_HOST and SMTP_USER and SMTP_PASSWORD)


def init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""CREATE TABLE IF NOT EXISTS users (
        username TEXT PRIMARY KEY,
        email TEXT UNIQUE,
        password_hash TEXT NOT NULL,
        salt TEXT NOT NULL,
        verified INTEGER DEFAULT 0,
        is_admin INTEGER DEFAULT 0
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS login_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT,
        timestamp TEXT,
        success INTEGER
    )""")
    conn.execute("""CREATE TABLE IF NOT EXISTS verification_codes (
        email TEXT,
        code TEXT,
        purpose TEXT,
        expires_at TEXT
    )""")
    conn.commit()
    conn.close()


def _hash_password(password: str, salt: str) -> str:
    return hashlib.sha256((salt + password).encode()).hexdigest()


def check_password_strength(password: str):
    """Returns (is_ok: bool, message: str)."""
    if len(password) < 8:
        return False, "Password must be at least 8 characters long."
    if not re.search(r"[A-Z]", password):
        return False, "Password must contain at least one UPPERCASE letter."
    if not re.search(r"[a-z]", password):
        return False, "Password must contain at least one lowercase letter."
    if not re.search(r"[0-9]", password):
        return False, "Password must contain at least one number."
    if not re.search(r"""[!@#$%^&*(),.?":{}|<>_\-+=\[\]/\\;'~`]""", password):
        return False, "Password must contain at least one special character (e.g. ! @ # $ %)."
    return True, "Strong password."


def send_email(to_email: str, subject: str, body: str) -> bool:
    """Sends an email via SMTP. If SMTP isn't configured, returns False so the
    caller can fall back to showing the code on-screen (useful for local testing)."""
    if not EMAIL_CONFIGURED:
        return False
    try:
        msg = MIMEText(body)
        msg["Subject"] = subject
        msg["From"] = SMTP_FROM
        msg["To"] = to_email
        context = ssl.create_default_context()
        with smtplib.SMTP(SMTP_HOST, int(SMTP_PORT)) as server:
            server.starttls(context=context)
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.sendmail(SMTP_FROM, to_email, msg.as_string())
        return True
    except Exception as e:
        print("Email send failed:", e)
        return False


def generate_code() -> str:
    return f"{secrets.randbelow(1000000):06d}"


def store_verification_code(email: str, purpose: str) -> str:
    code = generate_code()
    expires = (datetime.utcnow() + timedelta(minutes=15)).isoformat()
    conn = sqlite3.connect(DB_PATH)
    conn.execute("DELETE FROM verification_codes WHERE email = ? AND purpose = ?", (email, purpose))
    conn.execute("INSERT INTO verification_codes (email, code, purpose, expires_at) VALUES (?, ?, ?, ?)",
                 (email, code, purpose, expires))
    conn.commit()
    conn.close()
    return code


def check_verification_code(email: str, purpose: str, code: str) -> bool:
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT code, expires_at FROM verification_codes WHERE email = ? AND purpose = ?",
        (email, purpose),
    ).fetchone()
    if not row:
        conn.close()
        return False
    stored_code, expires_at = row
    valid = (stored_code == code.strip()) and (datetime.utcnow() <= datetime.fromisoformat(expires_at))
    if valid:
        conn.execute("DELETE FROM verification_codes WHERE email = ? AND purpose = ?", (email, purpose))
        conn.commit()
    conn.close()
    return valid


def create_user(username: str, email: str, password: str):
    salt = secrets.token_hex(16)
    ph = _hash_password(password, salt)
    is_admin = 1 if (ADMIN_USERNAME and username.strip().lower() == ADMIN_USERNAME.lower()) else 0
    conn = sqlite3.connect(DB_PATH)
    try:
        conn.execute(
            "INSERT INTO users (username, email, password_hash, salt, verified, is_admin) VALUES (?, ?, ?, ?, 0, ?)",
            (username.strip(), email.strip().lower(), ph, salt, is_admin),
        )
        conn.commit()
        return True, "Account created successfully."
    except sqlite3.IntegrityError:
        return False, "That username or email is already registered."
    finally:
        conn.close()


def mark_verified(email: str):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE users SET verified = 1 WHERE email = ?", (email,))
    conn.commit()
    conn.close()


def get_user_by_username(username: str):
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT username, email, password_hash, salt, verified, is_admin FROM users WHERE username = ?",
        (username.strip(),),
    ).fetchone()
    conn.close()
    return row


def update_password(email: str, new_password: str):
    salt = secrets.token_hex(16)
    ph = _hash_password(new_password, salt)
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE users SET password_hash = ?, salt = ? WHERE email = ?", (ph, salt, email))
    conn.commit()
    conn.close()


def log_login_attempt(username: str, success: bool):
    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        "INSERT INTO login_logs (username, timestamp, success) VALUES (?, ?, ?)",
        (username, datetime.utcnow().isoformat(), 1 if success else 0),
    )
    conn.commit()
    conn.close()


def verify_login(username: str, password: str):
    """Returns (ok: bool, message: str, is_admin: bool)."""
    row = get_user_by_username(username)
    if not row:
        log_login_attempt(username, False)
        return False, "Invalid username or password.", False
    _, email, stored_hash, salt, verified, is_admin = row
    if _hash_password(password, salt) != stored_hash:
        log_login_attempt(username, False)
        return False, "Invalid username or password.", False
    if not verified:
        log_login_attempt(username, False)
        return False, "Please verify your email before logging in (check the Verify Email page).", False
    log_login_attempt(username, True)
    return True, "Logged in successfully.", bool(is_admin)


def get_all_users():
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute("SELECT username, email, verified, is_admin FROM users ORDER BY username").fetchall()
    conn.close()
    return rows


def get_login_logs(limit: int = 200):
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT username, timestamp, success FROM login_logs ORDER BY id DESC LIMIT ?", (limit,)
    ).fetchall()
    conn.close()
    return rows


init_db()

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = None
if "is_admin" not in st.session_state:
    st.session_state.is_admin = False
if "pending_verify_email" not in st.session_state:
    st.session_state.pending_verify_email = None
if "pending_reset_email" not in st.session_state:
    st.session_state.pending_reset_email = None


# =====================================================================================
#  PAGES: Home / About / Sign Up / Verify Email / Login / Forgot Password / Admin
# =====================================================================================
def home_page():
    st.title(f"\U0001F4D8 Welcome to {APP_NAME}")
    st.markdown(f"""
A professional worksheet generator built for teachers, covering **Grades 1 to 10**.

**What you can do here:**
- Generate curriculum-based or free-content worksheets using AI (Anthropic Claude or Google Gemini)
- Or use your own question bank — completely free, no AI required
- Choose from Multiple Choice, Fill in the Blanks, Match the Column, Short, and Detailed question types
- Export polished, print-ready worksheets as PDF or Word, with your school's branding, logo, and layout

Use the menu on the left to **Sign Up** or **Log In**, then open the **Dashboard** to get started.
""")


def about_page():
    st.title("\u2139\ufe0f About SchoolSheet Studio")
    st.markdown("""
**SchoolSheet Studio** helps teachers create professional, ready-to-print worksheets in minutes.

**Features:**
- Grades 1-10 support with age-appropriate visual themes
- Multiple question types: MCQ, Fill in the Blanks, Match the Column, Short and Detailed Questions
- Three generation methods: curriculum-based, free-content, or your own question bank
- Proper Urdu and English language support
- PDF and Word export with custom branding, logos, and full layout control (spacing, binding gap, copies, randomization)
- Secure accounts with email verification and password reset

Built with a look and feel comparable to school platforms such as Cambridge, APS, and Roots-style institutions.
""")


def _password_field_with_toggle(label: str, key: str):
    """Renders a password text input with a 'Show password' checkbox next to it."""
    show = st.checkbox(f"Show {label.lower()}", key=f"{key}_show")
    return st.text_input(label, type="default" if show else "password", key=key)


def signup_page():
    st.title("\U0001F4DD Sign Up")
    if not EMAIL_CONFIGURED:
        st.warning(
            "Email sending isn't configured yet (no SMTP secrets set), so verification codes "
            "will be shown on-screen instead of emailed, for testing purposes."
        )
    with st.form("signup_form"):
        username = st.text_input("Username")
        email = st.text_input("Email")
        password = _password_field_with_toggle("Password", "signup_password")
        confirm_password = _password_field_with_toggle("Confirm Password", "signup_confirm_password")
        st.caption("Password must be 8+ characters and include an uppercase letter, a lowercase letter, a number, and a special character.")
        submitted = st.form_submit_button("Create Account")
    if submitted:
        if not username or not email or not password or not confirm_password:
            st.warning("Please fill in all fields.")
        elif password != confirm_password:
            st.error("Passwords do not match.")
        else:
            ok_strength, strength_msg = check_password_strength(password)
            if not ok_strength:
                st.error(strength_msg)
            else:
                ok, msg = create_user(username, email, password)
                if ok:
                    code = store_verification_code(email.strip().lower(), "signup")
                    emailed = send_email(
                        email, "Verify your SchoolSheet Studio account",
                        f"Your verification code is: {code}\nIt expires in 15 minutes."
                    )
                    st.session_state.pending_verify_email = email.strip().lower()
                    st.success(msg + " Go to the 'Verify Email' page to activate your account.")
                    if not emailed:
                        st.info(f"(Email not configured — your verification code is: **{code}**)")
                else:
                    st.error(msg)


def verify_email_page():
    st.title("\u2705 Verify Email")
    default_email = st.session_state.pending_verify_email or ""
    email = st.text_input("Email", value=default_email)
    code = st.text_input("Verification Code")
    if st.button("Verify"):
        if not email or not code:
            st.warning("Please enter both your email and the verification code.")
        elif check_verification_code(email.strip().lower(), "signup", code):
            mark_verified(email.strip().lower())
            st.success("Email verified successfully! You can now log in.")
            st.session_state.pending_verify_email = None
        else:
            st.error("Invalid or expired code. Please try again or sign up again to get a new code.")


def login_page():
    st.title("\U0001F511 Log In")
    username = st.text_input("Username")
    password = _password_field_with_toggle("Password", "login_password")
    if st.button("Log In"):
        if not username or not password:
            st.warning("Please enter both username and password.")
        else:
            ok, msg, is_admin = verify_login(username, password)
            if ok:
                st.session_state.logged_in = True
                st.session_state.username = username
                st.session_state.is_admin = is_admin
                st.success(msg)
                st.rerun()
            else:
                st.error(msg)
    st.caption("Forgot your password? Use the 'Forgot Password' page from the menu.")


def forgot_password_page():
    st.title("\U0001F513 Forgot Password")
    step = st.radio("Step", ["1. Request Code", "2. Reset Password"], horizontal=True)

    if step == "1. Request Code":
        email = st.text_input("Enter your account email")
        if st.button("Send Reset Code"):
            conn = sqlite3.connect(DB_PATH)
            row = conn.execute("SELECT username FROM users WHERE email = ?", (email.strip().lower(),)).fetchone()
            conn.close()
            if not row:
                st.error("No account found with that email.")
            else:
                code = store_verification_code(email.strip().lower(), "reset")
                emailed = send_email(
                    email, "SchoolSheet Studio password reset code",
                    f"Your password reset code is: {code}\nIt expires in 15 minutes."
                )
                st.session_state.pending_reset_email = email.strip().lower()
                st.success("Reset code generated. Go to Step 2 to enter it along with your new password.")
                if not emailed:
                    st.info(f"(Email not configured — your reset code is: **{code}**)")

    else:
        default_email = st.session_state.pending_reset_email or ""
        email = st.text_input("Email", value=default_email)
        code = st.text_input("Reset Code")
        new_password = _password_field_with_toggle("New Password", "reset_new_password")
        confirm_password = _password_field_with_toggle("Confirm New Password", "reset_confirm_password")
        if st.button("Reset Password"):
            if not email or not code or not new_password or not confirm_password:
                st.warning("Please fill in all fields.")
            elif new_password != confirm_password:
                st.error("Passwords do not match.")
            else:
                ok_strength, strength_msg = check_password_strength(new_password)
                if not ok_strength:
                    st.error(strength_msg)
                elif check_verification_code(email.strip().lower(), "reset", code):
                    update_password(email.strip().lower(), new_password)
                    st.success("Password reset successfully! You can now log in with your new password.")
                    st.session_state.pending_reset_email = None
                else:
                    st.error("Invalid or expired code.")


def admin_panel_page():
    st.title("\U0001F6E1\ufe0f Admin Panel")
    st.caption("Owner-only view: all registered users and full login history.")

    st.subheader("Registered Users")
    users = get_all_users()
    if users:
        st.dataframe(
            [{"Username": u, "Email": e, "Verified": bool(v), "Admin": bool(a)} for u, e, v, a in users],
            use_container_width=True,
        )
    else:
        st.info("No users registered yet.")

    st.subheader("Login History (most recent first)")
    logs = get_login_logs()
    if logs:
        st.dataframe(
            [{"Username": u, "Timestamp (UTC)": t, "Success": bool(s)} for u, t, s in logs],
            use_container_width=True,
        )
    else:
        st.info("No login attempts recorded yet.")


def render_dashboard():
    if "questions" not in st.session_state:
        st.session_state.questions = []

    st.title(f"\U0001F4D8 {APP_NAME}")
    st.caption("Professional Worksheet Generator — Grades 1 to 10")

    if AI_ENABLED:
        _provider_name = "Anthropic (Claude)" if ACTIVE_PROVIDER == "anthropic" else "Google (Gemini)"
        if st.session_state.is_admin:
            st.success(f"AI question generation is **enabled** using **{_provider_name}**.")
    else:
        if st.session_state.is_admin:
            st.warning(
                "AI question generation is **off** — no AI provider key is set. "
                "Add `ANTHROPIC_API_KEY` or `GEMINI_API_KEY` under Settings \u2192 Secrets to enable Methods 1 & 2 for all users. "
                "(Only you see this message, as the Owner/Admin.)"
            )
        else:
            st.info(
                "AI-assisted generation (Methods 1 & 2) isn't available right now — "
                "please use **Method 3: Your Own Question Bank** below instead."
            )

    with st.expander("\U0001F9ED Getting Started — Click here if you're new", expanded=True):
        st.markdown("""
    ### \U0001F44B New here? Start with this

    **1. Fill in Worksheet Settings below:** Choose Grade, Subject, Topic, Language, and Question Types — these apply to all 3 methods.

    **2. Generate questions — pick ONE of the 3 methods (tabs) below:**
    - **Method 1 — Curriculum-Based:** Upload a document (PDF/DOCX/TXT) containing content or questions to build the worksheet from it, OR simply generate directly from the Grade/Subject/Topic you selected above. *(Requires an AI provider)*
    - **Method 2 — Free Content:** Paste or upload any raw text (e.g. copied from ChatGPT) — it will be automatically cleaned and turned into a worksheet. *(Requires an AI provider)*
    - **Method 3 — Your Own Question Bank:** Upload a CSV/Excel file of pre-written questions — completely free, no AI required.

    **3. Go to the Preview & Export tab, choose PDF or Word format, and click "Create Worksheet" — a download button will appear immediately.**
    """)

    def _save_uploaded(uploaded_file):
        """Streamlit gives file-like objects, not paths — save to a temp path for the existing helpers."""
        if uploaded_file is None:
            return None
        suffix = os.path.splitext(uploaded_file.name)[1]
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=WORK_DIR)
        tmp.write(uploaded_file.getvalue())
        tmp.close()
        return tmp.name

    st.subheader("\u2699\ufe0f Worksheet Settings (apply to all methods)")

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        school_name = st.text_input("School Name", "Your School Name")
        grade = st.selectbox("Grade", GRADES, index=2)
    with c2:
        subject = st.selectbox("Subject", SUBJECTS, index=0)
        topic = st.text_input("Topic", placeholder="e.g. Nouns, Photosynthesis, Fractions")
    with c3:
        language = st.radio("Language", ["English", "Urdu"], horizontal=True)
        visual_theme = st.selectbox("Visual Theme", THEMES, index=0)
    with c4:
        cover_color = st.color_picker("Cover Color", "#2563EB")
        logo_file = st.file_uploader("School Logo (optional)", type=["png", "jpg", "jpeg"])

    q_types = st.multiselect("Question Types (select any combination)", QUESTION_TYPES, default=["Choose the Best Option (MCQ)"])

    c5, c6, c7, c8 = st.columns(4)
    with c5:
        num_q = st.number_input("Total Questions", min_value=1, value=10, step=1)
    with c6:
        q_per_page = st.number_input("Questions per Page (0 = unlimited)", min_value=0, value=0, step=1)
    with c7:
        num_copies = st.number_input("Number of Copies", min_value=1, value=1, step=1)
    with c8:
        copy_mode = st.radio("Copy Mode", ["Identical", "Randomized"], horizontal=True)

    c9, c10 = st.columns(2)
    with c9:
        spacing = st.slider("Spacing Between Questions (pt)", 0, 60, 18, step=2)
    with c10:
        binding_gap = st.slider("Binding Gap (mm)", 0, 40, 0, step=2)

    c11, c12, c13, c14 = st.columns(4)
    show_name = c11.checkbox("Show Student Name field", value=True)
    show_class = c12.checkbox("Show Class field", value=True)
    show_roll = c13.checkbox("Show Roll No. field", value=True)
    show_date = c14.checkbox("Show Date field", value=True)


    def build_cfg():
        logo_path = _save_uploaded(logo_file) if logo_file else None
        return WorksheetConfig(
            school_name=school_name or "Your School Name",
            logo_path=logo_path,
            logo_size_pt=60,
            cover_color_hex=cover_color,
            grade=str(grade),
            subject=subject,
            topic=topic or "General Topic",
            language=language,
            question_types=q_types or ["Choose the Best Option (MCQ)"],
            num_questions=int(num_q),
            questions_per_page=(int(q_per_page) if q_per_page and int(q_per_page) > 0 else None),
            spacing_pt=int(spacing),
            binding_gap_mm=int(binding_gap),
            num_copies=int(num_copies),
            copy_mode=copy_mode,
            visual_theme=visual_theme,
            show_student_name=show_name, show_class=show_class,
            show_roll_no=show_roll, show_date=show_date,
        )


    def preview_markdown(questions):
        if not questions:
            return "No questions generated yet."
        out = []
        for i, q in enumerate(questions, 1):
            out.append(f"**{i}. ({q.q_type})** {q.text}")
            if q.options:
                out.append("   " + "   ".join(f"({chr(97+j)}) {o}" for j, o in enumerate(q.options)))
            if q.match_left and q.match_right:
                out.append(f"   Match: {q.match_left}  \u2194  {q.match_right}")
            if q.image_keyword:
                out.append(f"   \U0001F5BC illustration: {q.image_keyword}")
            out.append("")
        return "\n".join(out)


    tab1, tab2, tab3, tab4 = st.tabs([
        "\U0001F4C2 Method 1: Curriculum-Based",
        "\U0001F4DD Method 2: Free Content",
        "\U0001F5C2\ufe0f Method 3: Question Bank",
        "\U0001F441\ufe0f Preview & Export",
    ])

    with tab1:
        st.markdown("Use **either** Path A or Path B.")
        pathA, pathB = st.tabs(["Path A — Upload File + Chapter/Topic/Page", "Path B — Use Grade/Subject/Topic above"])
        with pathA:
            m1a_file = st.file_uploader("Upload Curriculum File (PDF/DOCX/TXT)", key="m1a_file")
            colA1, colA2, colA3 = st.columns(3)
            m1a_chapter = colA1.text_input("Chapter", key="m1a_chapter")
            m1a_topic = colA2.text_input("Topic", key="m1a_topic")
            m1a_page = colA3.text_input("Page No.", key="m1a_page")
            if st.button("Generate from File \u2728", key="m1a_btn", type="primary"):
                if not m1a_file:
                    st.warning("Please upload a curriculum file first.")
                else:
                    try:
                        path = _save_uploaded(m1a_file)
                        raw = extract_text_from_file(path)
                        focus_note = f"\n\n[Focus strictly on: Chapter '{m1a_chapter}', Topic '{m1a_topic}', Page {m1a_page}]"
                        cfg = build_cfg()
                        cfg.topic = m1a_topic or cfg.topic
                        st.session_state.questions = generate_questions_ai(cfg, source_text=raw + focus_note)
                        st.success(f"Generated {len(st.session_state.questions)} questions.")
                    except RuntimeError as e:
                        st.error(str(e))
        with pathB:
            st.markdown("Uses the Grade / Subject / Topic selected above in Worksheet Settings.")
            if st.button("Generate from Selection \u2728", key="m1b_btn", type="primary"):
                try:
                    cfg = build_cfg()
                    st.session_state.questions = generate_questions_ai(cfg, source_text=None)
                    st.success(f"Generated {len(st.session_state.questions)} questions.")
                except RuntimeError as e:
                    st.error(str(e))

    with tab2:
        m2_text = st.text_area("Paste content here (e.g. from ChatGPT, notes, textbook excerpt)", height=200)
        m2_file = st.file_uploader("Or upload a file instead", key="m2_file")
        if st.button("Clean & Generate \u2728", key="m2_btn", type="primary"):
            raw = m2_text or ""
            if m2_file:
                raw += "\n" + extract_text_from_file(_save_uploaded(m2_file))
            if not raw.strip():
                st.warning("Please paste text or upload a file first.")
            else:
                try:
                    cleaned = clean_pasted_content(raw)
                    cfg = build_cfg()
                    st.session_state.questions = generate_questions_ai(cfg, source_text=cleaned)
                    st.success(f"Generated {len(st.session_state.questions)} questions.")
                except RuntimeError as e:
                    st.error(str(e))

    with tab3:
        st.markdown(
            "Upload a CSV or Excel file with columns: `grade, subject, topic, q_type, text, options, "
            "match_left, match_right, answer, image_keyword`. Use `|` to separate multiple options "
            "(e.g. `Cat|Dog|Cow|Horse`)."
        )
        m3_file = st.file_uploader("Upload Question Bank (CSV or XLSX)", key="m3_file")
        if st.button("Load & Build Worksheet \u2728", key="m3_btn", type="primary"):
            if not m3_file:
                st.warning("Please upload a question bank CSV/Excel file first.")
            else:
                try:
                    cfg = build_cfg()
                    df = load_question_bank(_save_uploaded(m3_file))
                    questions = questions_from_bank(df, cfg)
                    if not questions:
                        st.warning("No matching questions found for this Grade/Subject/Topic/Question-Type combination.")
                    else:
                        st.session_state.questions = questions
                        st.success(f"Loaded {len(questions)} questions.")
                except Exception as e:
                    st.error(f"Could not read question bank: {e}")

    with tab4:
        st.markdown(preview_markdown(st.session_state.questions))
        export_format = st.radio("Export Format", ["PDF", "Word"], horizontal=True)
        include_answer_key = st.checkbox("Include a separate Answer Key page at the end", value=True)
        if st.button("\U0001F4C4 Create Worksheet", type="primary"):
            if not st.session_state.questions:
                st.warning("Please generate/preview a worksheet first.")
            else:
                try:
                    cfg = build_cfg()
                    cfg.include_answer_key = include_answer_key
                    out_dir = tempfile.mkdtemp(dir=WORK_DIR)
                    if export_format == "PDF":
                        paths = build_pdf_copies(cfg, st.session_state.questions, out_dir)
                    else:
                        paths = build_docx_copies(cfg, st.session_state.questions, out_dir)

                    if len(paths) == 1:
                        with open(paths[0], "rb") as f:
                            st.download_button("\u2b07\ufe0f Download Worksheet", f, file_name=os.path.basename(paths[0]))
                        st.success(f"Worksheet generated ({export_format}).")
                    else:
                        zip_path = shutil.make_archive(os.path.join(out_dir, "worksheets"), "zip", out_dir)
                        with open(zip_path, "rb") as f:
                            st.download_button("\u2b07\ufe0f Download All Copies (.zip)", f, file_name="worksheets.zip")
                        st.success(f"{len(paths)} {export_format} copies generated and zipped.")
                except Exception as e:
                    st.error(f"Could not generate the worksheet file: {e}")
                    st.exception(e)


# =====================================================================================
#  SIDEBAR NAVIGATION
# =====================================================================================
with st.sidebar:
    st.markdown(f"## \U0001F4D8 {APP_NAME}")
    if st.session_state.logged_in:
        st.success(f"Logged in as **{st.session_state.username}**" + (" (Admin)" if st.session_state.is_admin else ""))
        nav_options = ["Home", "Dashboard", "About", "Logout"]
        if st.session_state.is_admin:
            nav_options.insert(1, "Admin Panel")
        nav = st.radio("Navigate", nav_options, label_visibility="collapsed")
    else:
        st.info("Log in or sign up to use the Dashboard.")
        nav = st.radio(
            "Navigate",
            ["Home", "About", "Login", "Sign Up", "Verify Email", "Forgot Password"],
            label_visibility="collapsed",
        )

if nav == "Home":
    home_page()
elif nav == "About":
    about_page()
elif nav == "Login":
    login_page()
elif nav == "Sign Up":
    signup_page()
elif nav == "Verify Email":
    verify_email_page()
elif nav == "Forgot Password":
    forgot_password_page()
elif nav == "Admin Panel":
    if st.session_state.is_admin:
        admin_panel_page()
    else:
        st.warning("Admin access only.")
elif nav == "Dashboard":
    if st.session_state.logged_in:
        render_dashboard()
    else:
        st.warning("Please log in to access the Dashboard.")
elif nav == "Logout":
    st.session_state.logged_in = False
    st.session_state.username = None
    st.session_state.is_admin = False
    st.success("You have been logged out.")
    st.rerun()
